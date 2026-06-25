from docx.oxml import parse_xml
from docx.oxml.shared import qn
from docx import Document

import html


def add_formatted_dropdown(paragraph, choices, default="", alias="", 
                         label="", style=None, instruction_text="Выберите значение"):
    """
    Добавляет форматированный выпадающий список с меткой
    
    Args:
        paragraph: параграф для добавления
        choices: список вариантов
        default: значение по умолчанию
        alias: уникальное имя
        label: текст метки перед списком
        style: стиль форматирования
        instruction_text: текст подсказки
    """
    # Добавляем метку, если она указана
    if label:
        run = paragraph.add_run(f"{label}: ")
        if style:
            run.style = style

    # Создаем XML для выпадающего списка
    sdt = parse_xml(f'''
        <w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:sdtPr>
                <w:alias w:val="{alias}"/>
                <w:tag w:val="{alias}"/>
                <w:id w:val="{abs(hash(alias))}"/>
                <w:dropDownList>
                    {''.join(f'<w:listItem w:displayText="{choice}" w:value="{choice}"/>' for choice in choices)}
                </w:dropDownList>
                <w:placeholder>
                    <w:docPart w:val="{instruction_text}"/>
                </w:placeholder>
            </w:sdtPr>
            <w:sdtContent>
                <w:r>
                    <w:rPr>
                        <w:color w:val="auto"/>
                        <w:sz w:val="24"/>
                    </w:rPr>
                    <w:t>{default if default else choices[0]}</w:t>
                </w:r>
            </w:sdtContent>
        </w:sdt>
    ''')
    
    paragraph._p.append(sdt)


def add_formatted_dropdown2(paragraph, choices, default="Не назначено", alias="", instruction_text=""):
    from xml.sax.saxutils import escape
    
    # Экранируем все строковые значения
    safe_alias = escape(str(alias))
    safe_default = escape(str(default))
    safe_instruction = escape(str(instruction_text))
    
    # Экранируем каждый choice
    safe_choices = []
    for choice in choices:
        if choice:
            # Экранируем специальные XML символы
            safe_choice = escape(str(choice))
            # Также экранируем кавычки для атрибутов
            safe_choice = safe_choice.replace('"', '&quot;')
            safe_choices.append(safe_choice)
    
    # Создаем XML элементы для выбора
    choices_xml = []
    for choice in safe_choices:
        choices_xml.append(f'<w:listItem w:displayText="{choice}" w:value="{choice}"/>')
    
    # Формируем XML
    dropdown_xml = f'''
        <w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:sdtPr>
                <w:alias w:val="{safe_alias}"/>
                <w:tag w:val="{safe_alias}"/>
                <w:id w:val="{abs(hash(safe_alias)) % 1000000}"/>
                <w:dropDownList>
                    <w:listItem w:displayText="{safe_default}" w:value="{safe_default}"/>
                    {''.join(choices_xml)}
                </w:dropDownList>
                <w:showingPlcHdr/>
                <w:placeholder>
                    <w:docPart w:val="{safe_instruction}"/>
                </w:placeholder>
            </w:sdtPr>
            <w:sdtContent>
                <w:r>
                    <w:rPr>
                        <w:color w:val="A0A0A0"/>
                        <w:sz w:val="24"/>
                        <w:spacing w:val="10"/>
                    </w:rPr>
                    <w:t>{safe_default}</w:t>
                </w:r>
            </w:sdtContent>
        </w:sdt>
    '''
    
    try:
        from docx.oxml import parse_xml
        sdt = parse_xml(dropdown_xml)
        paragraph._p.append(sdt)
    except Exception as e:
        # В случае ошибки вставляем текст
        print(f"Ошибка в add_formatted_dropdown2: {e}")
        paragraph.text = f"[{default}]"

def add_formatted_dropdown2_10pt(paragraph, choices, default="Не назначено", alias="", instruction_text=""):
    from xml.sax.saxutils import escape
    from docx.shared import Pt
    from docx.oxml.ns import qn

    # Экранируем строки
    safe_alias = escape(str(alias))
    safe_default = escape(str(default))
    safe_instruction = escape(str(instruction_text))
    
    safe_choices = []
    for choice in choices:
        if choice:
            safe_choice = escape(str(choice))
            safe_choice = safe_choice.replace('"', '&quot;')
            safe_choices.append(safe_choice)
    
    choices_xml = []
    for choice in safe_choices:
        choices_xml.append(f'<w:listItem w:displayText="{choice}" w:value="{choice}"/>')
    
    # Формируем XML для выпадающего списка
    # Основное изменение: добавляем элемент <w:doNotUsePlaceholder> и <w:rPr>
    dropdown_xml = f'''
        <w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:sdtPr>
                <w:alias w:val="{safe_alias}"/>
                <w:tag w:val="{safe_alias}"/>
                <w:id w:val="{abs(hash(safe_alias)) % 1000000}"/>
                <w:dropDownList>
                    <w:listItem w:displayText="{safe_default}" w:value="{safe_default}"/>
                    {''.join(choices_xml)}
                </w:dropDownList>
                <!-- ВАЖНО: Отключаем плейсхолдер -->
                <w:showingPlcHdr/> 
                <!-- ВАЖНО: Задаем форматирование для вводимого текста -->
                <w:rPr>
                    <w:sz w:val="20"/> <!-- 10pt = 20 полупунктов -->
                    <w:szCs w:val="20"/>
                </w:rPr>
            </w:sdtPr>
            <w:sdtContent>
                <w:r>
                    <w:rPr>
                        <w:color w:val="A0A0A0"/>                    
                        <w:sz w:val="20"/>
                        <w:szCs w:val="20"/>
                    </w:rPr>
                    <w:t>{safe_default}</w:t>
                </w:r>
            </w:sdtContent>
        </w:sdt>
    '''
    
    try:
        from docx.oxml import parse_xml
        sdt = parse_xml(dropdown_xml)
        paragraph._element.append(sdt)
    except Exception as e:
        print(f"Ошибка в add_formatted_dropdown2_10pt: {e}")
        # fallback
        run = paragraph.add_run(f"[{default}]")
        run.font.size = Pt(10)

def add_formatted_dropdown3(paragraph, inputs_choices, controls_choices=[], default="Не назначено", alias="", instruction_text="", first_divider='Сигналы РЗиА', second_divider='Общие сигналы ФС'):
    # Экранируем все переменные с помощью html.escape
    default_esc = html.escape(str(default))
    alias_esc = html.escape(str(alias))
    instruction_esc = html.escape(str(instruction_text))
    first_divider_esc = html.escape(str(first_divider))
    second_divider_esc = html.escape(str(second_divider))
    
    # Формируем элементы списка с разделителями
    list_items = []
    
    # Добавляем inputs
    if inputs_choices:
        list_items.append(f'<w:listItem w:displayText="──────── {first_divider_esc} ────────" w:value="INPUTS_HEADER" w:disabled="true"/>')
        for choice in inputs_choices:
            choice_esc = html.escape(str(choice))
            list_items.append(f'<w:listItem w:displayText="{choice_esc}" w:value="{choice_esc}"/>')
    
    # Добавляем controls
    if controls_choices:
        if inputs_choices:  # Добавляем разделитель только если есть оба списка
            list_items.append('<w:listItem w:displayText=" " w:value="SPACER" w:disabled="true"/>')
        list_items.append(f'<w:listItem w:displayText="────── {second_divider_esc} ──────" w:value="CONTROLS_HEADER" w:disabled="true"/>')
        for choice in controls_choices:
            choice_esc = html.escape(str(choice))
            list_items.append(f'<w:listItem w:displayText="{choice_esc}" w:value="{choice_esc}"/>')
    
    sdt = parse_xml(f'''
        <w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:sdtPr>
                <w:alias w:val="{alias_esc}"/>
                <w:tag w:val="{alias_esc}"/>
                <w:id w:val="{abs(hash(alias))}"/>
                <w:dropDownList>
                    <w:listItem w:displayText="{default_esc}" w:value="{default_esc}"/>
                    {''.join(list_items)}
                </w:dropDownList>
                <w:showingPlcHdr/>
                <w:placeholder>
                    <w:docPart w:val="{instruction_esc}"/>
                </w:placeholder>
            </w:sdtPr>
            <w:sdtContent>
                <w:r>
                    <w:rPr>
                        <w:color w:val="A0A0A0"/>
                        <w:sz w:val="24"/>
                        <w:spacing w:val="10"/>
                    </w:rPr>
                    <w:t>{default_esc}</w:t>
                </w:r>
            </w:sdtContent>
        </w:sdt>
    ''')
    
    paragraph._p.append(sdt)










if __name__ == '__main__':
    # Пример использования расширенной версии:
    doc = Document()

    # Добавляем выпадающий список с меткой и форматированием
    paragraph = doc.add_paragraph()
    choices = ["Высокий", "Средний", "Низкий"]
    add_formatted_dropdown(
        paragraph=paragraph,
        choices=choices,
        default="Выберите уровень",
        alias="PriorityLevel",
        label="Приоритет",
        instruction_text="Выберите уровень приоритета"
    )

    # Добавляем несколько выпадающих списков в таблицу
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'

    # Заполняем таблицу выпадающими списками
    for i, (label, choices) in enumerate([
        ("Статус", ["Активен", "Неактивен", "В обработке"]),
        ("Категория", ["A", "B", "C"]),
        ("Важность", ["Критическая", "Высокая", "Средняя", "Низкая"])
    ]):
        cell = table.cell(i, 0)
        cell.text = label
        
        cell = table.cell(i, 1)
        paragraph = cell.paragraphs[0]
        add_formatted_dropdown(
            paragraph=paragraph,
            choices=choices,
            #alias= f"DropDown_{i}",
            instruction_text=f"Выберите {label.lower()}",
        )

    doc.save("formatted_dropdowns.docx")
