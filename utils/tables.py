
from docx.shared import Cm, Inches
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt
import docx
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from docx.enum.table import WD_ROW_HEIGHT_RULE


import os, sys
import json

from pathlib import Path

from .dropdowns import add_formatted_dropdown2, add_formatted_dropdown3, add_formatted_dropdown2_10pt


def set_table_borders(table):
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')  # 8 = 1pt (значение в восьмых долях пункта)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    tblPr.append(tblBorders)


def set_vertical_cell_direction(cell: _Cell, direction: str):
    # direction: tbRl -- top to bottom, btLr -- bottom to top
    assert direction in ("tbRl", "btLr")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)  # btLr tbRl
    tcPr.append(textDirection)

def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

def set_cell_vertical_alignment(cell, align="center"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')
        tcValign.set(qn('w:val'), align)
        tcPr.append(tcValign)

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

####################################################################################
################################ ТАБЛИЦА ДЛЯ УСТАВОК ХАЛЕЗОВ 07-08-25 ##############################
####################################################################################

table_settings = (Inches(0.25), Inches(1.6), Inches(1.1), Inches(1.6), Inches(0.45), Inches(0.45), Inches(1.4), Inches(1), Inches(1), Inches(1), Inches(1))  #задаем ширину столбцов таблицы вывода репортов

def add_table_settings(doc):
    table = doc.add_table(rows=5, cols=11)
    table.style = 'Сетка таблицы51'
    table.allow_autofit = False
    set_table_borders(table)

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Наименование'
    hdr_cells[3].text = 'Значение / Диапазон'
    hdr_cells[4].text = 'Ед. изм.'
    hdr_cells[5].text = 'Шаг'   
    hdr_cells[6].text = 'Значение по умолчанию'
    hdr_cells[7].text = 'Группы уставок'

    for i in range(0,10):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице

    # p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # p.runs[0].font.size = Pt(10)

    hdr_cells = table.rows[1].cells # вторая строка заголовка таблицы
    hdr_cells[1].text = 'ПО ЮС'
    hdr_cells[2].text = 'ИЧМ'
    hdr_cells[7].text = '1'
    hdr_cells[8].text = '2'
    hdr_cells[9].text = '3'
    hdr_cells[10].text = '4'    
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    hdr_cells[7].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[8].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[9].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[10].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # третья строка со служебными тегами
    hdr_cells = table.rows[2].cells
    #hdr_cells[2].text = '{%tr for param_name, param_data in input_value.properties.items() %}'
    tag = f'for row in func.settings'
    hdr_cells[2].text = '{%tr '+ tag + ' %}'
    # четвертая строка со служебными тегами
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '{{ loop.index }}'
    hdr_cells[1].text = '{{ row[0] }}'
    hdr_cells[2].text = '{{ row[1] }}'    
    hdr_cells[3].text = '{{ row[2]  }}'
    hdr_cells[4].text = '{{ row[3] }}'
    hdr_cells[5].text = '{{ row[4] }}'
    hdr_cells[6].text = '{{ row[5] }}'
    #hdr_cells[8].text = '' #'{{ param_data.setpoint }}'

 
    hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #hdr_cells[7].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # пятая строка со служебными тегами
    hdr_cells = table.rows[4].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,9):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек
    table.cell(0, 1).merge(table.cell(0, 2))
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 3).merge(table.cell(1, 3))
    table.cell(0, 4).merge(table.cell(1, 4))
    table.cell(0, 5).merge(table.cell(1, 5))
    table.cell(0, 6).merge(table.cell(1, 6))
    table.cell(0, 7).merge(table.cell(0, 10))

    table.cell(2, 0).merge(table.cell(2, 10))
    table.cell(4, 0).merge(table.cell(4, 10))

    for row in table.rows:
        for idx, width in enumerate(table_settings):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить

        # Устанавливаем высоту шрифта (11 пунктов) для всех ячеек таблицы
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Устанавливаем размер шрифта 12 пунктов

    return table

####################################################################################
################################ КОНЕЦ ТАБЛИЦА ДЛЯ УСТАВОК #########################
####################################################################################



####################################################################################
############################ ТАБЛИЦА ДЛЯ МАТРИЦЫ ДИСКРЕТНЫХ ВХОДОВ ###############
####################################################################################

table_mtrx_ins = (Inches(2), Inches(4))

def add_table_mtrx_ins(doc, inputs, controls=[]): # новая таблица исходящих отчетов
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Дискретный вход'
    hdr_cells[1].text = 'Назначенный сигнал'

    for i in range(0,2):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице

    hdr_cells = table.rows[1].cells
    tag = r'for i in range(1, input_module[1]|int+1)'
    hdr_cells[0].text = '{%tr '+ tag + ' %}'

    # четвертая строка со служебными тегами
    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = 'Дискретный вход '+'{{ loop.index }}'


    #choices_start = ["Не выполняется", "По переднему фронту", "По заднему фронту", "По любому изменению"]
    par1 = hdr_cells[1].paragraphs[0]
    add_formatted_dropdown3(
        paragraph=par1,
        inputs_choices=inputs,
        controls_choices=controls,
    )
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # пятая строка со служебными тегами
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,2):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек

    table.cell(1, 0).merge(table.cell(1, 1))
    table.cell(3, 0).merge(table.cell(3, 1))

    for row in table.rows:
        for idx, width in enumerate(table_mtrx_ins):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить
    return table 

####################################################################################
############################ КОНЕЦ ТАБЛИЦА ДЛЯ МАТРИЦЫ ДИСКРЕТНЫХ ВХОДОВ ###########
####################################################################################


####################################################################################
############################ ТАБЛИЦА ДЛЯ МАТРИЦЫ ВЫХОДНЫХ РЕЛЕ #####################
####################################################################################

table_mtrx_outs = (Inches(2), Inches(1.7), Inches(1.7), Inches(1.7), Inches(1.7), Inches(1.7))

def add_table_mtrx_outs(doc, statuses, controls=[]): # новая таблица исходящих отчетов
    table = doc.add_table(rows=5, cols=6)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Выходное реле'
    hdr_cells[1].text = 'Назначенные сигналы'

    for i in range(0,6):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице

    hdr_cells = table.rows[1].cells # вторая строка заголовка таблицы
    hdr_cells[1].text = '1'
    hdr_cells[2].text = '2'
    hdr_cells[3].text = '3'
    hdr_cells[4].text = '4'
    hdr_cells[5].text = '5'    
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    hdr_cells = table.rows[2].cells
    tag = r'for i in range(1, output_module[1]|int+1)'
    hdr_cells[2].text = '{%tr '+ tag + ' %}'

    # четвертая строка со служебными тегами
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = 'Реле '+'{{ loop.index }}'


    #choices_start = ["Не выполняется", "По переднему фронту", "По заднему фронту", "По любому изменению"]
    par1 = hdr_cells[1].paragraphs[0]
    add_formatted_dropdown3(
        paragraph=par1,
        inputs_choices=statuses,
        controls_choices=controls,
        #alias= f"DropDown_{i}",
        #instruction_text=f"Выберите ",
    )
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    par2 = hdr_cells[2].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par2,
        choices=statuses,
    )
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    par3 = hdr_cells[3].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par3,
        choices=statuses,
    )
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    par4 = hdr_cells[4].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par4,
        choices=statuses,
    )
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    par5 = hdr_cells[5].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par5,
        choices=statuses,
    )
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # пятая строка со служебными тегами
    hdr_cells = table.rows[4].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,6):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(0, 5))

    table.cell(2, 0).merge(table.cell(2, 4))
    table.cell(4, 0).merge(table.cell(4, 4))

    for row in table.rows:
        for idx, width in enumerate(table_mtrx_outs):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить
    return table 

####################################################################################
############################ КОНЕЦ ТАБЛИЦА ДЛЯ МАТРИЦЫ ВЫХОДНЫХ РЕЛЕ ###############
####################################################################################


####################################################################################
######################## ТАБЛИЦА ДЛЯ СВЕТОДИОДОВ УСОВЕРШЕНСТВОВАННАЯ ###############
####################################################################################

table_leds_new = (Inches(1.7), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5))

def add_table_leds_new(doc, statuses, plates_data): # новая таблица исходящих отчетов
    table = doc.add_table(rows=5, cols=7)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Светодиод'
    hdr_cells[1].text = 'Режим работы'
    hdr_cells[2].text = 'Назначенный сигнал 1'    
    hdr_cells[3].text = 'Назначенный сигнал 2' 
    hdr_cells[4].text = 'Назначенный сигнал 3' 
    hdr_cells[5].text = 'Назначенный сигнал 4' 
    hdr_cells[6].text = 'Назначенный сигнал 5' 

    for i in range(0,7):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице

    hdr_cells = table.rows[1].cells
    tag = f'for i in range(1, 17)'
    hdr_cells[0].text = '{%tr '+ tag + ' %}'

    # четвертая строка со служебными тегами
    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = 'Светодиод '+'{{ loop.index }}' + ' (красный)'
    # четвертая строка со служебными тегами
    hdr_cells_row2 = table.rows[3].cells
    hdr_cells_row2[0].text = 'Светодиод '+'{{ loop.index }}'  + ' (зеленый)'

    choices = ["С фиксацией"]
    par2 = hdr_cells[1].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par2,
        choices=choices,
        default='Без фиксации')

    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    par1 = hdr_cells[2].paragraphs[0]
    add_formatted_dropdown3(
        paragraph=par1,
        inputs_choices=statuses,
        controls_choices = plates_data,
        first_divider= 'Сигналы РЗиА',
        second_divider= 'Сигналы от блоков')
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    par21 = hdr_cells_row2[1].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par21,
        choices=choices,
        default='Без фиксации')
    hdr_cells_row2[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 

    par12 = hdr_cells_row2[2].paragraphs[0]
    add_formatted_dropdown3(
        paragraph=par12,
        inputs_choices=statuses,
        controls_choices = plates_data,
        first_divider= 'Сигналы РЗиА',
        second_divider= 'Сигналы от блоков')
    hdr_cells_row2[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # первая строка
    par22 = hdr_cells[3].paragraphs[0]
    add_formatted_dropdown3(
        paragraph=par22,
        inputs_choices=statuses,
        controls_choices = plates_data,
        first_divider= 'Сигналы РЗиА',
        second_divider= 'Сигналы от блоков')
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
    par23 = hdr_cells[4].paragraphs[0]

    add_formatted_dropdown3(
        paragraph=par23,
        inputs_choices=statuses,
        controls_choices = plates_data,
        first_divider= 'Сигналы РЗиА',
        second_divider= 'Сигналы от блоков')
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
    par24 = hdr_cells[5].paragraphs[0]

    add_formatted_dropdown3(
        paragraph=par24,
        inputs_choices=statuses,
        controls_choices = plates_data,
        first_divider= 'Сигналы РЗиА',
        second_divider= 'Сигналы от блоков')
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
    par25 = hdr_cells[6].paragraphs[0]

    add_formatted_dropdown3(
        paragraph=par25,
        inputs_choices=statuses,
        controls_choices = plates_data,
        first_divider= 'Сигналы РЗиА',
        second_divider= 'Сигналы от блоков')
    hdr_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 


    # вторая строка заполнение со второго сигнала
    par13 = hdr_cells_row2[3].paragraphs[0]
    add_formatted_dropdown3(
        paragraph=par13,
        inputs_choices=statuses,
        controls_choices = plates_data,
        first_divider= 'Сигналы РЗиА',
        second_divider= 'Сигналы от блоков',)
    hdr_cells_row2[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    par14 = hdr_cells_row2[4].paragraphs[0]

    add_formatted_dropdown3(
        paragraph=par14,
        inputs_choices=statuses,
        controls_choices = plates_data,
        first_divider= 'Сигналы РЗиА',
        second_divider= 'Сигналы от блоков')
    hdr_cells_row2[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    par15 = hdr_cells_row2[5].paragraphs[0]

    add_formatted_dropdown3(
        paragraph=par15,
        inputs_choices=statuses,
        controls_choices = plates_data,
        first_divider= 'Сигналы РЗиА',
        second_divider= 'Сигналы от блоков')
    hdr_cells_row2[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    par16 = hdr_cells_row2[6].paragraphs[0]
    add_formatted_dropdown3(
        paragraph=par16,
        inputs_choices=statuses,
        controls_choices = plates_data,
        first_divider= 'Сигналы РЗиА',
        second_divider= 'Сигналы от блоков')
    hdr_cells_row2[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # пятая строка со служебными тегами
    hdr_cells = table.rows[4].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,3):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек

    table.cell(1, 0).merge(table.cell(1, 2))
    table.cell(4, 0).merge(table.cell(4, 2))

    for row in table.rows:
        for idx, width in enumerate(table_leds_new):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить
    return table 

####################################################################################
################## КОНЕЦ ТАБЛИЦА ДЛЯ СВЕТОДИОДОВ УСОВЕРШЕНСТВОВАННАЯ ###############
####################################################################################


####################################################################################
############################ ТАБЛИЦА ДЛЯ ФУНКЦИОНАЛЬНЫХ КЛАВИШ ###############
####################################################################################

table_fks = (Inches(2), Inches(4))

def add_table_fks(doc, choices_start): # новая таблица исходящих отчетов
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Функциональная клавиша'
    hdr_cells[1].text = 'Назначенный сигнал'

    for i in range(0,2):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице

    hdr_cells = table.rows[1].cells
    tag = f'for i in range(1, 17)'
    hdr_cells[0].text = '{%tr '+ tag + ' %}'

    # четвертая строка со служебными тегами
    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = 'Функциональная клавиша '+'{{ loop.index }}'

    par1 = hdr_cells[1].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par1,
        choices=choices_start,
    )
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # пятая строка со служебными тегами
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,2):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек

    table.cell(1, 0).merge(table.cell(1, 1))
    table.cell(3, 0).merge(table.cell(3, 1))

    for row in table.rows:
        for idx, width in enumerate(table_fks):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить
    return table

####################################################################################
############################ КОНЕЦ ТАБЛИЦА ДЛЯ ФУНКЦИОНАЛЬНЫХ КЛАВИШ ###############
####################################################################################


#########################################  НОВАЯ  ###################################
################################ ТАБЛИЦА ДЛЯ ДИСКРЕТНЫХ ВХОДОВ ВЫХОДОВ  #############
#####################################################################################

table_binaries = (Inches(0.28), Inches(1.23), Inches(1.4), Inches(1.5), Inches(0.55), Inches(0.45), Inches(0.9), Inches(1.05))  #задаем ширину столбцов таблицы вывода репортов

def add_table_binaries(doc, tag = 'for row in items'):
    table = doc.add_table(rows=4, cols=8)
    table.style = 'Сетка таблицы51'
    table.allow_autofit = False
    set_table_borders(table)

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Описание'
    hdr_cells[2].text = 'Наименование'
    hdr_cells[3].text = 'Значение / Диапазон'
    hdr_cells[4].text = 'Ед. изм.'
    hdr_cells[5].text = 'Шаг'
    hdr_cells[6].text = 'Значение по умолчанию'
    hdr_cells[7].text = 'Уставка'
    for i in range(0,8):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице

    # p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # p.runs[0].font.size = Pt(10)

    #hdr_cells = table.rows[1].cells # вторая строка заголовка таблицы
    #hdr_cells[2].text = 'ПО'
    #hdr_cells[3].text = 'ФСУ'
    #hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # третья строка со служебными тегами
    hdr_cells = table.rows[1].cells
    #hdr_cells[2].text = '{%tr for param_name, param_data in input_value.properties.items() %}'
    #tag = f'for row in items'
    hdr_cells[2].text = '{%tr '+ tag + ' %}'
    # четвертая строка со служебными тегами
    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = '{{ loop.index }}'
    hdr_cells[1].text = '{{ row[0] }}'
    hdr_cells[2].text = '{{ row[1] }}'
    #hdr_cells[3].text = '{{ row["Наименование ФСУ"] }}'    
    hdr_cells[3].text = '{{ row[2]  }}'
    hdr_cells[4].text = '{{ row[3] }}'
    hdr_cells[5].text = '{{ row[4] }}'
    hdr_cells[6].text = '{{ row[5] }}'
    hdr_cells[7].text = '' #'{{ param_data.setpoint }}'

    hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[7].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # пятая строка со служебными тегами
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,8):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек
    #table.cell(0, 2).merge(table.cell(0, 3))
    #table.cell(0, 0).merge(table.cell(1, 0))
    #table.cell(0, 1).merge(table.cell(1, 1))
    #table.cell(0, 4).merge(table.cell(1, 4))
    #table.cell(0, 5).merge(table.cell(1, 5))
    #table.cell(0, 6).merge(table.cell(1, 6))
    #table.cell(0, 7).merge(table.cell(1, 7))
    #table.cell(0, 8).merge(table.cell(1, 8))

    table.cell(1, 0).merge(table.cell(1, 7))
    table.cell(3, 0).merge(table.cell(3, 7))

    for row in table.rows:
        for idx, width in enumerate(table_binaries):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить

        # Устанавливаем высоту шрифта (11 пунктов) для всех ячеек таблицы
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)  # Устанавливаем размер шрифта 12 пунктов

    return table

####################################################################################
######## КОНЕЦ ТАБЛИЦА ДЛЯ ДИСКРЕТНЫХ ФХОДОВ ВЫХОДОВ НОВАЯ #########################
####################################################################################


####################################################################################
################################ ТАБЛИЦА ДЛЯ РЕГИСТРАЦИИ ###########################
####################################################################################


table_reg = (Inches(4.5), Inches(1.5), Inches(1.6), Inches(1.6), Inches(1.6))  #задаем ширину столбцов таблицы вывода репортов

def add_table_reg(doc, tag = 'for row in fsu.get_statuses()'): # новая таблица исходящих отчетов
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Параметр'
    hdr_cells[2].text = 'Журнал событий регистрация'
    hdr_cells[3].text = 'Осциллограф пуск'
    hdr_cells[4].text = 'Осциллограф регистрация'
    for i in range(0,5):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице


    hdr_cells = table.rows[1].cells # вторая строка заголовка таблицы
    hdr_cells[0].text = 'Наименование'
    hdr_cells[1].text = 'Обозначение ФСУ'
    hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    hdr_cells = table.rows[2].cells
      
    hdr_cells[2].text = '{%tr '+ tag + ' %}'

    # четвертая строка со служебными тегами
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '{{ row[0] }}'
    hdr_cells[1].text = '{{ row[1] }}'

    #hdr_cells[2].text = '{{ param_data.log }}'
    choices_start = ["Не выполняется", "По переднему фронту", "По заднему фронту", "По любому изменению"]
    par3 = hdr_cells[2].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par3,
        choices=choices_start,
        #alias= f"DropDown_{i}",
        #instruction_text=f"Выберите ",
    )
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #hdr_cells[3].text = '{{ param_data.oscill_start }}'
    par2 = hdr_cells[3].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par2,
        choices=choices_start,
        #alias= f"DropDown_{i}",
        #instruction_text=f"Выберите ",
    )
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #hdr_cells[4].text = '+'
    choices_reg = ["Выведено", "Введено"]
    par1 = hdr_cells[4].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par1,
        choices=choices_reg,
        #alias= f"DropDown_{i}",
        #instruction_text=f"Выберите ",
    )
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # пятая строка со служебными тегами
    hdr_cells = table.rows[4].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,5):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 2).merge(table.cell(1, 2))
    table.cell(0, 3).merge(table.cell(1, 3))
    table.cell(0, 4).merge(table.cell(1, 4))

    table.cell(2, 0).merge(table.cell(2, 4))
    table.cell(4, 0).merge(table.cell(4, 4))

    for row in table.rows:
        for idx, width in enumerate(table_reg):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить

            # Устанавливаем высоту шрифта (11 пунктов) для всех ячеек таблицы
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)  # Устанавливаем размер шрифта 12 пунктов

    return table    

####################################################################################
################################ КОНЕЦ ТАБЛИЦА ДЛЯ РЕГИСТРАЦИИ #####################
####################################################################################

####################################################################################
############################ ФИНАЛЬНАЯ ТАБЛИЦА С ПОДПИСЯМИ СОСТАВИТЕЛЯ ###############
####################################################################################

table_final = (Inches(3), Inches(3))

def add_table_final(doc): # новая таблица исходящих отчетов
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Стиль5'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ФИО составителя:'

    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = 'Номер и дата составления:'

    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = 'Дата выдачи:'

    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = 'Дата окончания:'


    table.allow_autofit = False
    table.autofit = False
    table.style = 'Стиль5'

    # --- Форматирование таблицы ---
    for row in table.rows:
        for idx, width in enumerate(table_final):
            row.cells[idx].width = width
            # Установка высоты строки (новый код)
            row.height = Pt(20)  # Укажите нужную высоту в пунктах
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY  # Фиксированная высота

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)


    return table 



table_settings_core4 = (Inches(0.25), Inches(1.6), Inches(1.1), Inches(1.6), Inches(0.45), Inches(0.45), Inches(1.4), Inches(1), Inches(1), Inches(1), Inches(1)) 

def add_table_settings_core4(doc):
    """
    Создает таблицу уставок для Core4
    """
    table = doc.add_table(rows=2, cols=11)
    table.style = 'Сетка таблицы51'
    table.allow_autofit = False
    set_table_borders(table)
    
    # Устанавливаем фиксированный макет таблицы
    tbl_pr = table._tbl.xpath('./w:tblPr')[0]
    # Удаляем старый tblLayout если есть, чтобы избежать дублирования
    for elem in tbl_pr.xpath('./w:tblLayout'):
        tbl_pr.remove(elem)
    
    tbl_pr.append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )
    
    # --- ЗАПОЛНЕНИЕ ЗАГОЛОВКОВ ---
    
    # ПЕРВАЯ СТРОКА ЗАГОЛОВКА
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Наименование'
    hdr_cells[3].text = 'Значение / Диапазон'
    hdr_cells[4].text = 'Ед. изм.'
    hdr_cells[5].text = 'Шаг'   
    hdr_cells[6].text = 'Значение по умолчанию'
    hdr_cells[7].text = 'Группы уставок'


    # Настраиваем стиль и выравнивание
    # Индексы, которые должны быть по ЦЕНТРУ: 
    # 0 (№), 1 (Наименование), 3 (Значение), 4 (Ед.изм), 5 (Шаг), 6 (По умолч.), 7 (Группы)

    for i in range(0,10):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    set_repeat_table_header(table.rows[0])

    # ВТОРАЯ СТРОКА ЗАГОЛОВКА (подзаголовки)
    hdr_cells_2 = table.rows[1].cells
    hdr_cells_2[1].text = 'ПО ЮС'
    hdr_cells_2[2].text = 'ИЧМ'
    hdr_cells_2[7].text = '1'
    hdr_cells_2[8].text = '2'
    hdr_cells_2[9].text = '3'
    hdr_cells_2[10].text = '4'
    
    # Выравнивание для второй строки
    for idx in [1, 2, 3, 7, 8, 9, 10]:
        hdr_cells_2[idx].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- СЛИЯНИЕ ЯЧЕЕК ---
    # Важно: порядок слияния может иметь значение. Обычно лучше сливать сверху вниз или слева направо.
    table.cell(0, 1).merge(table.cell(0, 2))       # Наименование (охватывает кол 1 и 2)
    table.cell(0, 0).merge(table.cell(1, 0))       # № (вертикальное слияние)
    table.cell(0, 3).merge(table.cell(1, 3))       # Значение (вертикальное)
    table.cell(0, 4).merge(table.cell(1, 4))       # Ед. изм. (вертикальное)
    table.cell(0, 5).merge(table.cell(1, 5))       # Шаг (вертикальное)
    table.cell(0, 6).merge(table.cell(1, 6))       # По умолчанию (вертикальное)
    table.cell(0, 7).merge(table.cell(0, 10))      # Группы уставок (горизонтальное слияние 7-10)


    # --- УСТАНОВКА ШИРИНЫ СТОЛБЦОВ (КЛЮЧЕВОЕ ИЗМЕНЕНИЕ) ---
    # Вместо установки ширины каждой ячейке, устанавливаем ширину каждому столбцу.
    # Это работает корректно даже при фиксированном макете и слияниях.
    for col_idx, width in enumerate(table_settings_core4):
        # Получаем объект столбца
        column = table.columns[col_idx]
        # Устанавливаем ширину столбца
        column.width = width
        
        # Дополнительно можно продублировать установку ширины для первой ячейки столбца,
        # так как некоторые версии Word/библиотек лучше реагируют на cell.width
        # Но column.width является основным драйвером при fixed layout.
        try:
            table.cell(0, col_idx).width = width
        except IndexError:
            pass # Если ячейка была удалена из-за слияния, пропускаем


    # --- УСТАНОВКА ШРИФТА ---
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table



# Ширина столбцов для режима 10 колонок (сокращенный вид)
table_settings_core4_simple = (
    Inches(0.25),  # №
    Inches(1.5),   # Наименование (объединенное)
    Inches(1.8),   # Значение / Диапазон
    Inches(0.5),   # Ед. изм.
    Inches(0.5),   # Шаг
    Inches(1.5),   # По умолчанию
    Inches(1),     # Группа 1
    Inches(1),     # Группа 2
    Inches(1),     # Группа 3
    Inches(1)      # Группа 4
) 

def add_table_settings_core4_simple(doc):
    """
    Создает таблицу уставок для Core4 (режим 1 - сокращенный).
    Убирает разделение ПО ЮС/ИЧМ, оставляя одно общее "Наименование".
    Сохраняет блок "Группы уставок".
    """
    # Создаем таблицу из 10 колонок
    table = doc.add_table(rows=2, cols=10)
    table.style = 'Сетка таблицы51'
    table.allow_autofit = False
    
    # Фиксированный макет
    tbl_pr = table._tbl.xpath('./w:tblPr')[0]
    for elem in tbl_pr.xpath('./w:tblLayout'):
        tbl_pr.remove(elem)
    tbl_pr.append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )
    
    # --- ПЕРВАЯ СТРОКА ЗАГОЛОВКА ---
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Наименование'
    hdr_cells[2].text = 'Значение / Диапазон'
    hdr_cells[3].text = 'Ед. изм.'
    hdr_cells[4].text = 'Шаг'   
    hdr_cells[5].text = 'Значение по умолчанию'
    hdr_cells[6].text = 'Группы уставок'

    # Стиль и выравнивание для первой строки
    for i in range(0, 7): # Заголовки до индекса 6 включительно
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        try:
            set_cell_vertical_alignment(hdr_cells[i], align="center")
        except:
            pass
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    set_repeat_table_header(table.rows[0])

    # --- ВТОРАЯ СТРОКА ЗАГОЛОВКА (подзаголовки) ---
    hdr_cells_2 = table.rows[1].cells
    # Подзаголовки для групп уставок
    hdr_cells_2[6].text = '1'
    hdr_cells_2[7].text = '2'
    hdr_cells_2[8].text = '3'
    hdr_cells_2[9].text = '4'
    
    # Выравнивание для второй строки
    for idx in [6, 7, 8, 9]:
        hdr_cells_2[idx].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- СЛИЯНИЕ ЯЧЕЕК ---
    # 1. Наименование (объединяем ячейку 1 и 2? Нет, в этом режиме у нас одна колонка Наименования)
    # Но так как мы убрали ПО ЮС и ИЧМ, нам не нужно горизонтальное слияние для Наименования.
    # Однако, нам нужно вертикальное слияние для остальных колонок, чтобы они выглядели красиво.
    
    table.cell(0, 0).merge(table.cell(1, 0))       # № (вертикальное)
    table.cell(0, 1).merge(table.cell(1, 1))       # Наименование (вертикальное)
    table.cell(0, 2).merge(table.cell(1, 2))       # Значение (вертикальное)
    table.cell(0, 3).merge(table.cell(1, 3))       # Ед. изм. (вертикальное)
    table.cell(0, 4).merge(table.cell(1, 4))       # Шаг (вертикальное)
    table.cell(0, 5).merge(table.cell(1, 5))       # По умолчанию (вертикальное)
    table.cell(0, 6).merge(table.cell(0, 9))       # Группы уставок (горизонтальное слияние 6-9)

    # --- УСТАНОВКА ШИРИНЫ СТОЛБЦОВ ---
    for col_idx, width in enumerate(table_settings_core4_simple):
        try:
            table.columns[col_idx].width = width
            table.cell(0, col_idx).width = width
        except IndexError:
            pass

    # --- ШРИФТ ---
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table




####################################################################################
############################ ТАБЛИЦА ДЛЯ МАТРИЦЫ ДИСКРЕТНЫХ ВХОДОВ ###############
####################################################################################


# Константа ширины колонок
TABLE_WIDTHS_MTRX_INS_CORE4 = (Inches(2), Inches(4))

def add_table_mtrx_ins_core4(doc, slot_name, inputs_list, sigs, di_sigs):
    """
    Создает статическую таблицу для одного слота без использования Jinja2.
    
    :param doc: Объект Document
    :param slot_name: Имя слота (для заголовка или контекста)
    :param inputs_list: Список описаний входов (например, ['Слот M8. ДВ1', ...])
    :param sigs: Список доступных сигналов для dropdown (из get_fsu_signals)
    :param di_sigs: Список дискретных сигналов (если нужно разделить логику)
    :return: Объект Table
    """
    
    if not inputs_list:
        return None

    # Создаем таблицу: 1 строка заголовка + N строк данных
    num_rows = len(inputs_list) + 1
    table = doc.add_table(rows=num_rows, cols=2)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # Фиксируем макет таблицы
    try:
        tbl_pr = table._tbl.xpath('./w:tblPr')[0]
        tbl_layout = parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
        tbl_pr.append(tbl_layout)
    except Exception:
        pass # Игнорируем ошибки настройки XML, если таблица простая

    # --- СТРОКА 0: ЗАГОЛОВКИ ---
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Дискретный вход'
    hdr_cells[1].text = 'Назначенный сигнал'

    for i in range(2):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Если функция доступна
        if 'set_cell_vertical_alignment' in globals():
            set_cell_vertical_alignment(hdr_cells[i], align="center")

    # Настройка повторения заголовка (если поддерживается)
    if 'set_repeat_table_header' in globals():
        set_repeat_table_header(table.rows[0])

    # --- ЗАПОЛНЕНИЕ ДАННЫХ ---
    for idx, input_desc in enumerate(inputs_list):
        row_idx = idx + 1
        row_cells = table.rows[row_idx].cells
        
        # Левая ячейка: Описание входа
        row_cells[0].text = input_desc
        p_left = row_cells[0].paragraphs[0]
        p_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT # Или CENTER, по вкусу
        # Можно задать стиль для обычных ячеек, если нужно
        # p_left.style = 'Normal'

        # Правая ячейка: Выпадающий список / Элемент управления
        par_right = row_cells[1].paragraphs[0]
        
        # Вызываем вашу функцию добавления dropdown
        # Предполагаем, что она модифицирует paragraph на месте
        if 'add_formatted_dropdown3' in globals():
            add_formatted_dropdown3(
                paragraph=par_right,
                inputs_choices=sigs,      # Передаем общие сигналы
                controls_choices=di_sigs, # Передаем дискретные сигналы
            )
        else:
            # Заглушка, если функции нет
            par_right.text = "[Нет сигнала]"
            
        par_right.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- НАСТРОЙКА ШИРИНЫ КОЛОНОК ---
    for row in table.rows:
        try:
            row.cells[0].width = TABLE_WIDTHS_MTRX_INS_CORE4[0]
            row.cells[1].width = TABLE_WIDTHS_MTRX_INS_CORE4[1]
        except Exception:
            pass

    return table




# Ширина колонок для таблицы выходов: [Имя реле, Сигнал 1, Сигнал 2, ..., Сигнал 5]
TABLE_WIDTHS_OUTS = (Inches(2), Inches(1.7), Inches(1.7), Inches(1.7), Inches(1.7), Inches(1.7))

def add_table_mtrx_outs_core4(doc, outputs_list, sigs_list):
    """
    Создает статическую таблицу параметрирования выходных реле.
    Каждая строка - одно реле.
    Колонки: Имя реле + 5 колонок с dropdown для выбора сигналов.
    """
    if not outputs_list:
        return
    
    num_rows = len(outputs_list) + 2  # 2 строки заголовка + данные
    table = doc.add_table(rows=num_rows, cols=6)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # Фиксируем макет
    try:
        tbl_pr = table._tbl.xpath('./w:tblPr')[0]
        tbl_layout = parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
        tbl_pr.append(tbl_layout)
    except Exception:
        pass

    # --- СТРОКА 0: ГЛАВНЫЕ ЗАГОЛОВКИ ---
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Выходное реле'
    hdr_cells[1].text = 'Назначенные сигналы'

    for i in range(0, 6):
        p = hdr_cells[i].paragraphs[0]
        try: p.style = 'ДОК Таблица Заголовок'
        except: pass
        if 'set_cell_vertical_alignment' in globals():
            set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- СТРОКА 1: ПОДЗАГОЛОВКИ (номера 1-5) ---
    hdr_cells2 = table.rows[1].cells
    hdr_cells2[1].text = '1'
    hdr_cells2[2].text = '2'
    hdr_cells2[3].text = '3'
    hdr_cells2[4].text = '4'
    hdr_cells2[5].text = '5'
    
    for i in range(1, 6):
        p = hdr_cells2[i].paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        try: p.style = 'ДОК Таблица Заголовок'
        except: pass
        if 'set_cell_vertical_alignment' in globals():
            set_cell_vertical_alignment(hdr_cells2[i], align="center")

    # --- ОБЪЕДИНЕНИЕ ЯЧЕЕК ---
    # Объединяем "Выходное реле" вертикально (строка 0 и строка 1)
    table.cell(0, 0).merge(table.cell(1, 0))
    
    # Объединяем "Назначенные сигналы" горизонтально (колонки 1-5 в строке 0)
    table.cell(0, 1).merge(table.cell(0, 5))

    # --- ЗАПОЛНЕНИЕ ДАННЫХ ---
    for idx, output_name in enumerate(outputs_list):
        row_idx = idx + 2
        row_cells = table.rows[row_idx].cells
        
        # Колонка 0: Имя реле
        row_cells[0].text = output_name
        row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Колонки 1-5: Выпадающие списки
        for col_idx in range(1, 6):
            par = row_cells[col_idx].paragraphs[0]
            par.clear()
            
            if 'add_formatted_dropdown2' in globals():
                add_formatted_dropdown2(
                    paragraph=par,
                    choices=sigs_list,
                )
            else:
                par.text = ""
            
            par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- ПОВТОРЕНИЕ ЗАГОЛОВКОВ НА СЛЕДУЮЩИХ СТРАНИЦАХ ---
    if 'set_repeat_table_header' in globals():
        set_repeat_table_header(table.rows[0])  # Первая строка
        set_repeat_table_header(table.rows[1])  # Вторая строка (как в старом коде)

    # --- ПРИМЕНЕНИЕ ШИРИНЫ КОЛОНОК ---
    for row in table.rows:
        for idx, width in enumerate(TABLE_WIDTHS_OUTS):
            if idx < len(row.cells):
                row.cells[idx].width = width

    return table




####################################################################################
######################## ТАБЛИЦА ДЛЯ СВЕТОДИОДОВ УСОВЕРШЕНСТВОВАННАЯ ###############
####################################################################################


# Define widths for 8 columns: 
# [LED Name, Mode, Color, Sig1, Sig2, Sig3, Sig4, Sig5]
TABLE_WIDTHS_LEDS = (
    Inches(1.2),  # LED Name
    Inches(1.2),  # Mode
    Inches(1.0),  # Color
    Inches(1.5),  # Sig 1
    Inches(1.5),  # Sig 2
    Inches(1.5),  # Sig 3
    Inches(1.5),  # Sig 4
    Inches(1.5)   # Sig 5
)

def add_table_leds_new_core4(doc, statuses, led_count=16):
    """
    Creates a table for LED configuration with dynamic rows.
    
    Args:
        doc: python-docx Document object.
        statuses: List of choices for signal dropdowns.
        plates_data: List of choices for signal dropdowns (controls).
        led_count: Number of LEDs (rows) to generate (default 16).
    """
    
    # 1. Create Table: Header row + led_count data rows
    total_rows = 1 + led_count
    table = doc.add_table(rows=total_rows, cols=8)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # 2. Set Fixed Layout
    tbl_pr = table._tbl.xpath('./w:tblPr')
    if tbl_pr:
        tbl_pr[0].append(
            parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
        )

    # 3. Configure Header Row (Row 0)
    hdr_cells = table.rows[0].cells
    headers = [
        'Светодиод', 
        'Режим работы', 
        'Цвет', 
        'Назначенный сигнал 1',    
        'Назначенный сигнал 2', 
        'Назначенный сигнал 3', 
        'Назначенный сигнал 4', 
        'Назначенный сигнал 5'
    ]
    
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Assuming set_cell_vertical_alignment is defined in your context
        try:
            set_cell_vertical_alignment(hdr_cells[i], align="center")
        except NameError:
            pass 


    # Set Header Repeat (if table spans multiple pages)
    set_repeat_table_header(table.rows[0]) 

    # 4. Generate Data Rows (LED 1 to LED 16)
    for row_idx in range(1, total_rows):
        row = table.rows[row_idx]
        cells = row.cells
        
        # Column 0: LED Name (e.g., "Светодиод 1")
        cells[0].text = f'Светодиод {row_idx}'
        p_name = cells[0].paragraphs[0]
        p_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Column 1: Mode Dropdown
        par_mode_data = cells[1].paragraphs[0]
        add_formatted_dropdown2(
            paragraph=par_mode_data,
            choices=["С фиксацией"],
            default='Без фиксации'
        )
        par_mode_data.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Column 2: Color Dropdown
        par_color_data = cells[2].paragraphs[0]
        add_formatted_dropdown2(
            paragraph=par_color_data,
            choices=['Зеленый'],
            default='Красный'
        )
        par_color_data.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Columns 3-7: Signal Dropdowns
        for col_idx in range(3, 8):
            par_sig_data = cells[col_idx].paragraphs[0]
            add_formatted_dropdown2(
                paragraph=par_sig_data,
                choices=statuses
            )
            par_sig_data.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 5. Apply Column Widths
    for row in table.rows:
        for idx, width in enumerate(TABLE_WIDTHS_LEDS):
            if idx < len(row.cells):
                row.cells[idx].width = width

    return table



####################################################################################
############################ ТАБЛИЦА ДЛЯ ФУНКЦИОНАЛЬНЫХ КЛАВИШ ###############
####################################################################################

# Widths for 2 columns: [Key Name, Signal Dropdown]
TABLE_WIDTHS_KFS = (Inches(2), Inches(4))

def add_table_fks_core4(doc, choices, key_count=16):
    """
    Creates a table for Functional Keys (FKs) configuration.
    
    Args:
        doc: python-docx Document object.
        choices: List of choices for the signal dropdown.
        key_count: Number of functional keys (rows) to generate (default 16).
    """
    
    # 1. Create Table: 1 Header row + key_count data rows
    total_rows = 1 + key_count
    table = doc.add_table(rows=total_rows, cols=2)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # 2. Set Fixed Layout
    tbl_pr = table._tbl.xpath('./w:tblPr')
    if tbl_pr:
        tbl_pr[0].append(
            parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
        )

    # 3. Configure Header Row (Row 0)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Функциональная клавиша'
    hdr_cells[1].text = 'Назначенный сигнал'

    for i in range(2):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        try:
            set_cell_vertical_alignment(hdr_cells[i], align="center")
        except NameError:
            pass

    # Set Header Repeat (if table spans multiple pages)
    set_repeat_table_header(table.rows[0])

    # 4. Generate Data Rows (FK 1 to FK 16)
    for row_idx in range(1, total_rows):
        row = table.rows[row_idx]
        cells = row.cells
        
        # Column 0: Key Name (e.g., "Функциональная клавиша 1")
        cells[0].text = f'Функциональная клавиша {row_idx}'
        p_name = cells[0].paragraphs[0]
        p_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Column 1: Signal Dropdown
        par_sig = cells[1].paragraphs[0]
        add_formatted_dropdown2(
            paragraph=par_sig,
            choices=choices
        )
        par_sig.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 5. Apply Column Widths
    for row in table.rows:
        for idx, width in enumerate(TABLE_WIDTHS_KFS):
            if idx < len(row.cells):
                row.cells[idx].width = width

    return table




#########################################  НОВАЯ  ###################################
################################ ТАБЛИЦА ДЛЯ ДИСКРЕТНЫХ ВХОДОВ ВЫХОДОВ  #############
#####################################################################################

table_binaries4 = (Inches(0.28), Inches(1.23), Inches(1.4), Inches(1.5), Inches(0.55), Inches(0.45), Inches(0.9), Inches(1.05))  # задаем ширину столбцов таблицы вывода репортов

def add_table_binaries_core4(doc, data_rows):
    """
    Создает таблицу с данными параметров
    
    Args:
        doc: документ python-docx
        data_rows: список кортежей с данными строк (col1, col2, col3, col4, col5, col6)
    """
    if not data_rows:
        return None
    
    # Создаем таблицу: заголовок + тело (на 1 строку больше, чем данных)
    table = doc.add_table(rows=1 + len(data_rows), cols=8)
    table.style = 'Сетка таблицы51'
    table.allow_autofit = False
    set_table_borders(table)

    # Устанавливаем фиксированный макет таблицы
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )
    
    # Заполняем заголовок таблицы (первая строка)
    hdr_cells = table.rows[0].cells
    headers = ['№', 'Описание', 'Обозначение ФСУ', 'Значение / Диапазон', 
               'Ед. изм.', 'Шаг', 'Значение по умолчанию', 'Уставка']
    
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Повторяем заголовок на следующей странице
    set_repeat_table_header(table.rows[0])

    # Заполняем строки данных
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        
        # Номер строки (начиная с 1)
        row_cells[0].text = str(row_idx + 1)
        row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Данные строки
        row_cells[1].text = str(row_data[0]) if row_data[0] else ''  # Описание
        row_cells[2].text = str(row_data[1]) if row_data[1] else ''  # Обозначение ФСУ
        row_cells[3].text = str(row_data[2]) if row_data[2] else ''  # Значение / Диапазон
        row_cells[4].text = str(row_data[3]) if row_data[3] else ''  # Ед. изм.
        row_cells[5].text = str(row_data[4]) if row_data[4] else ''  # Шаг
        row_cells[6].text = str(row_data[5]) if row_data[5] else ''  # Значение по умолчанию
        row_cells[7].text = ''  # Уставка (пустое поле для заполнения)
        
        # Выравнивание для числовых полей
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[7].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Устанавливаем ширину столбцов
    for row in table.rows:
        for idx, width in enumerate(table_binaries4):
            if idx < len(row.cells):
                row.cells[idx].width = width

    # Устанавливаем высоту шрифта (11 пунктов) для всех ячеек таблицы
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    return table



####################################################################################
################################ ТАБЛИЦА ДЛЯ РЕГИСТРАЦИИ ###########################
####################################################################################


table_reg4 = (Inches(4.0), Inches(2.0), Inches(1.6), Inches(1.6), Inches(1.6))

def add_table_reg_core4(doc, data_rows):
    """
    Создает таблицу для настройки параметров регистрации с выпадающими списками
    
    Args:
        doc: документ python-docx
        data_rows: список кортежей (наименование, обозначение ФСУ, restrain)
    """
    if not data_rows:
        return None
    
    # Создаем таблицу: 2 строки заголовка + строки данных
    table = doc.add_table(rows=2 + len(data_rows), cols=5)
    table.style = 'Стиль7'
    table.allow_autofit = False
    
    # Устанавливаем фиксированный макет таблицы
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )
    
    # ====== ПЕРВАЯ СТРОКА ЗАГОЛОВКА ======
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Параметр'
    hdr_cells[2].text = 'Журнал событий регистрация'
    hdr_cells[3].text = 'Осциллограф пуск'
    hdr_cells[4].text = 'Осциллограф регистрация'
    for i in range(5):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    
    # ====== ВТОРАЯ СТРОКА ЗАГОЛОВКА ======
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = 'Наименование'
    hdr_cells[1].text = 'Обозначение ФСУ'
    hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # ====== ОБЪЕДИНЕНИЕ ЯЧЕЕК ======
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 2).merge(table.cell(1, 2))
    table.cell(0, 3).merge(table.cell(1, 3))
    table.cell(0, 4).merge(table.cell(1, 4))
    
    # ====== ВАРИАНТЫ ДЛЯ ВЫПАДАЮЩИХ СПИСКОВ ======
    choices_reg = ["Введено"]
    choices_osc = ["По переднему фронту", "По заднему фронту", "По любому изменению"]
    
    # ====== ЗАПОЛНЕНИЕ ДАННЫХ С ВЫПАДАЮЩИМИ СПИСКАМИ ======
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 2].cells
        
        col1, col2, type = row_data
        
        # Столбец 0: Наименование (обычный текст)
        row_cells[0].text = str(col1) if col1 else ''
        row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Столбец 1: Обозначение ФСУ (обычный текст)
        row_cells[1].text = str(col2) if col2 else ''
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Столбец 2: Журнал событий регистрация (выпадающий список)
        add_formatted_dropdown2_10pt(
            paragraph=row_cells[2].paragraphs[0],
            choices=choices_osc,
            default="Не выполняется"
        )
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Столбец 3: Осциллограф пуск (выпадающий список)
        if type != 3:
            row_cells[3].text = 'Не выполняется'
        else:
            add_formatted_dropdown2_10pt(
                paragraph=row_cells[3].paragraphs[0],
                choices=choices_osc,
                default="Не выполняется"
            )        
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Столбец 4: Осциллограф регистрация (выпадающий список)
        if type != 3:
            row_cells[4].text = 'Выведено'
        else:        
            add_formatted_dropdown2_10pt(
                paragraph=row_cells[4].paragraphs[0],
                choices=choices_reg,
                default="Выведено"
            )
        row_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # ====== НАСТРОЙКА ШИРИНЫ СТОЛБЦОВ ======
    for row in table.rows:
        for idx, width in enumerate(table_reg4):
            if idx < len(row.cells):
                row.cells[idx].width = width
    
    # ====== УСТАНОВКА РАЗМЕРА ШРИФТА ======
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    set_repeat_table_header(table.rows[0])
    set_repeat_table_header(table.rows[1])

    return table