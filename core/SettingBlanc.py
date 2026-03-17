# Класс, представляющий бланк уставок
# Требует инициализации
# для генерации бланка уставок в него нужно передать объект класса Device

from docxtpl import DocxTemplate
from docx import Document

from utils.docx_handler import add_new_section, add_new_section_landscape
from utils.tables import add_table_settings, add_table_mtrx_ins, add_table_mtrx_outs, add_table_leds_new, add_table_fks, add_table_binaries, add_table_reg, add_table_final

from xml.sax.saxutils import escape # для экранирования в дропдаун списке всяких << >>

from docxtpl import DocxTemplate

from logger.logger import Logger

class SettingBlanc:
    def __init__(self, code = '', versions = [{"edition":"X.X", "data": "XX.XX.XXXX"},]):
        self.code = code
        self.versions = versions

    # СУММАРНАЯ ФУНКЦИЯ СОЗДАНИЯ РАЗДЕЛА УСТАВКИ РЗиА
    def _create_section_settings(self, fsu, doc):
        if not fsu.has_any_setting():
            return
        add_new_section_landscape(doc)

        p = doc.add_paragraph('УСТАВКИ РЗиА'+r'{% for fb in fsu.merged_fbs %}')
        p.style = 'ДОК Заголовок 1'

        p = doc.add_paragraph(r'{{ fb.description_fb }} ({{ fb.russian_name }}) {% for func in fb.functions %}')
        #p.style = 'ДОК Заголовок 3'
        p.style = 'ДОК Заголовок 2'

        p = doc.add_paragraph(r'{{ func.description }} ({{ func.name }})')
        p.style = 'ДОК Таблица Название'

        add_table_settings(doc)

        p = doc.add_paragraph(r'{% endfor %}{% endfor %}')
        p.style = 'TAGS'

    # РАЗДЕЛ ПАРАМЕТРИРОВАНИЯ ВХОДОВ И ВЫХОДОВ
    def _create_section_inouts(self, modules, fsu, doc):
        if not modules.get_modules_binary_inputs() and not modules.get_modules_binary_outputs():
            return

        add_new_section_landscape(doc) # Создаем раздел для матрицы вх/вых
        # Добавляем заголовок
        p = doc.add_paragraph('МАТРИЦА ВХОДОВ И ВЫХОДНЫХ РЕЛЕ')
        p.style = 'ДОК Заголовок 1'

        if modules.get_modules_binary_inputs():
            p = doc.add_paragraph('Дискретные входы')
            p.style = 'ДОК Заголовок 2'
            text = doc.add_paragraph('Для дискретного входа возможно подключение только одного сигнала.')
            text.style = 'ДОК Текст'
            p = doc.add_paragraph(r'{% for input_module in modules.get_modules_binary_inputs() %}')
            p.style = 'ДОК Текст'
            p = doc.add_paragraph(r'{{ input_module[0] }}')
            p.style = 'ДОК Таблица Название'
            inputs = fsu.get_inputs()
            #inputs = sorted(list(inputs))
            inputs = sorted([item[0] for item in inputs])
            controls = fsu.get_controls()
            #controls = sorted(list(controls))
            controls = sorted([item[0] for item in controls])
            add_table_mtrx_ins(doc, inputs, controls)
            p = doc.add_paragraph(r'{% endfor %}')
            p.style = 'TAGS'   

        ############################### ВЫХОДНЫЕ РЕЛЕ ####################################
        if modules.get_modules_binary_outputs():
            p = doc.add_paragraph('Выходные реле')
            p.style = 'ДОК Заголовок 2'

            text = doc.add_paragraph('Возможно подключение до пяти сигналов на одно выходное реле.')
            text.style = 'ДОК Текст'

            p = doc.add_paragraph(r'{% for output_module in modules.get_modules_binary_outputs() %}')
            p.style = 'ДОК Текст'

            p = doc.add_paragraph(r'{{ output_module[0] }}')
            p.style = 'ДОК Таблица Название'

            statuses = fsu.get_statuses()
            #statuses = sorted(list(statuses))
            statuses = sorted([item[0] for item in statuses])

            add_table_mtrx_outs(doc, statuses)

            p = doc.add_paragraph(r'{% endfor %}')
            p.style = 'TAGS' 

    # РАЗДЕЛ СВЕТОДИОДОВ И ФК
    def _create_section_leds(self, modules, hmi, fsu, doc):
        if hmi.order_code=='': # Если ИЧМ не заказан, но раздел не формируем
            return

        #############################################################################
        # СОЗДАЕМ РАЗДЕЛ НАСТРОЙКА СВЕТОДИОДОВ И ФУНКЦИОНАЛЬНЫХ КЛАВИШ
        add_new_section_landscape(doc) # Создаем раздел для матрицы вх/вых
        # Добавляем заголовок
        p = doc.add_paragraph('НАСТРОЙКА СВЕТОДИОДОВ И ФУНКЦИОНАЛЬНЫХ КЛАВИШ')
        p.style = 'ДОК Заголовок 1'

        ###############################################################
        p = doc.add_paragraph('Светодиоды')
        p.style = 'ДОК Заголовок 2'

        text = doc.add_paragraph('Для светодиода возможно подключение до пяти сигналов.')
        text.style = 'ДОК Текст'

        p = doc.add_paragraph(r'{% for leds in hmi.get_leds() if hmi.get_leds() %}')
        p.style = 'ДОК Текст'

        p = doc.add_paragraph(r'{{ leds }}')
        p.style = 'ДОК Таблица Название'

        statuses = fsu.get_statuses()
        statuses = sorted([item[0] for item in statuses])

        add_table_leds_new(doc, statuses, plates_data=modules.get_statuses())

        p = doc.add_paragraph(r'{% endfor %}')
        p.style = 'TAGS'    

        ###############################################################
        p = doc.add_paragraph('Функциональные клавиши')
        p.style = 'ДОК Заголовок 2'

        text = doc.add_paragraph('Для функциональной клавиши возможно подключение только одного управляющего сигнала.')
        text.style = 'ДОК Текст'

        p = doc.add_paragraph(r'{% for fks in hmi.get_fks() if hmi.get_fks() %}')
        p.style = 'ДОК Текст'

        p = doc.add_paragraph(r'{{ fks }}')
        p.style = 'ДОК Таблица Название'

        choices = fsu.get_controls()
        #choices = sorted(list(choices))
        choices = sorted([item[0] for item in choices])

        add_table_fks(doc, choices)

        p = doc.add_paragraph(r'{% endfor %}')
        p.style = 'TAGS'

    # СУММАРНАЯ ФУНКЦИЯ СОЗДАНИЯ РАЗДЕЛА КОНФИГУРАЦИЯ
    def _create_section_config(self, aux_funcs, doc):

        ############################################################################
        # СОЗДАЕМ РАЗДЕЛ С КОНФИГУРАЦИЕЙ

        if not aux_funcs.get_config_sync() and not aux_funcs.get_config_cpu() and not aux_funcs.get_config_disturb():
            return

        add_new_section(doc)
        p = doc.add_paragraph('КОНФИГУРАЦИЯ')
        p.style = 'ДОК Заголовок 1'

        #### СИНХРОНИЗАЦИЯ ВРЕМЕНИ ####
        if aux_funcs.get_config_sync():    
            p = doc.add_paragraph('Синхронизация времени')
            p.style = 'ДОК Заголовок 2'

            p = doc.add_paragraph(r'Общие настройки синхронизации{% set items = aux_funcs.get_config_sync()["Общие параметры"] %}')
            p.style = 'ДОК Таблица Название'
            add_table_binaries(doc)

            p = doc.add_paragraph()
            p.style = 'TAGS'

            p = doc.add_paragraph(r'Параметры летнего времени{% set items = aux_funcs.get_config_sync()["Параметры летнего времени"] %}')
            p.style = 'ДОК Таблица Название'
            add_table_binaries(doc)

            p = doc.add_paragraph()
            p.style = 'TAGS'

            p = doc.add_paragraph(r'Параметры зимнего времени{% set items = aux_funcs.get_config_sync()["Параметры зимнего времени"] %}')
            p.style = 'ДОК Таблица Название'
            add_table_binaries(doc)

        #### МОДУЛЬ ЦП ####
        if aux_funcs.get_config_cpu():
            p = doc.add_paragraph('Модуль ЦП')
            p.style = 'ДОК Заголовок 2'

            p = doc.add_paragraph(r'Резервирование{% set items = aux_funcs.get_config_cpu()["Резервирование"] %}')
            p.style = 'ДОК Таблица Название'
            add_table_binaries(doc)

            p = doc.add_paragraph()
            p.style = 'TAGS'

        #### НАСТРОЙКА РЕГИСТРАЦИИ ####
        if aux_funcs.get_config_disturb():
            p = doc.add_paragraph('Настройка регистрации')
            p.style = 'ДОК Заголовок 2'

            p = doc.add_paragraph(r'Общие параметры{% set items = aux_funcs.get_config_disturb()["Общие уставки"] %}')
            p.style = 'ДОК Таблица Название'
            add_table_binaries(doc)

            p = doc.add_paragraph()
            p.style = 'TAGS'

        # СОЗДАЕМ РАЗДЕЛ С КОНФИГУРАЦИЕЙ ВХОДОВ ВЫХОДОВ
        #p = doc.add_paragraph('Модули'+r'{% for plate in hardware.get_hw_plates() if hardware.get_hw_plates() %}')
        #p.style = 'ДОК Заголовок 2'

        p = doc.add_paragraph(r'{% for module in modules if modules %}')
        p.style = 'TAGS'

        p = doc.add_paragraph(r'Слот М{{ module.slot_number }}. {{ module.description }} ({{ module.russian_name }})'+ r'{% set items = module.obj.get_settings_for_configuration_docx_table()[“Общие настройки конфигурации”] %}' + r'{% if items %}')
        #p.style = 'ДОК Заголовок 3'
        p.style = 'ДОК Заголовок 2'

        p = doc.add_paragraph(r'Общие настройки конфигурации'+r'{% set item = items[0] %}')
        p.style = 'ДОК Таблица Название'
        add_table_binaries(doc, tag='for row in item.settings')
        p = doc.add_paragraph(r'{% endif %}{% set items = module.obj.get_settings_for_configuration_docx_table()["Входы"] %}{% if items %}{% for item in items %}')
        p.style = 'TAGS'

        p1 = doc.add_paragraph(r'{{ item.description }}')
        p1.style = 'ДОК Таблица Название'
        add_table_binaries(doc, tag='for row in item.settings')
        p = doc.add_paragraph(r'{% endfor %}{% endif %}{% set items = module.obj.get_settings_for_configuration_docx_table()[“Выходы”] %}{% if items %}{% for item in items %}')
        p.style = 'TAGS'

        p = doc.add_paragraph(r'{{ item.description }}')
        p.style = 'ДОК Таблица Название'
        add_table_binaries(doc, tag='for row in item.settings')
        p = doc.add_paragraph(r'{% endfor %}{% endif %}{% endfor %}')
        p.style = 'TAGS'

    # РАЗДЕЛ НАСТРОЙКИ ПАРАМЕТРОВ РЕГИСТРАЦИИ
    def _create_section_disturb(self, fsu, doc):
        if not fsu.get_statuses() and not fsu.get_controls() and not fsu.get_inputs():
            return
        add_new_section_landscape(doc)

        p = doc.add_paragraph('НАСТРОЙКА ПАРАМЕТРОВ РЕГИСТРАЦИИ')
        p.style = 'ДОК Заголовок 1'

        text = doc.add_paragraph('Возможна регистрация не более 200 сигналов.')
        text.style = 'ДОК Текст'

        # Таблица Выходные сигналы общей логики
        #if fsu.get_fsu_sys_statuses_sorted():
            #p = doc.add_paragraph('Системные сигналы')
            #p.style = 'ДОК Таблица Название'
            #add_table_reg(doc, generate = 0)

        # Таблица Выходные сигналы общей логики
        if fsu.get_statuses():
            p = doc.add_paragraph('Выходные сигналы общей логики')
            p.style = 'ДОК Таблица Название'
            add_table_reg(doc, tag = 'for row in fsu.get_statuses()') # 0 - сист сигналы, 1 - статусы, 2 - вирт ключи и кнопки, 3 - входные сигналы
            space = doc.add_paragraph(' ')
            space.style = 'TAGS'

        # Таблица Виртуальные ключи и клавиши
        if fsu.get_controls():
            p = doc.add_paragraph('Виртуальные ключи и клавиши')
            p.style = 'ДОК Таблица Название'
            add_table_reg(doc, tag = 'for row in fsu.get_controls()')
            space = doc.add_paragraph(' ')
            space.style = 'TAGS'

        # Таблица Входные дискретные сигналы
        if fsu.get_inputs():    
            p = doc.add_paragraph('Входные дискретные сигналы')
            p.style = 'ДОК Таблица Название'
            add_table_reg(doc, tag = 'for row in fsu.get_inputs()')
            space = doc.add_paragraph(' ')
            space.style = 'TAGS'        

    def create_template(self, fsu, modules, hmi, aux_funcs):
        doc = Document('origin.docx')
        # СОЗДАЕМ РАЗДЕЛ С УСТАВКАМИ РЗИА
        self._create_section_settings(fsu, doc)
        # СОЗДАЕМ РАЗДЕЛ С ПАРАМЕТРИРОВАНИЕ ДИСКРЕТНЫХ ВХОДОВ И ВЫХОДНЫХ РЕЛЕ
        self._create_section_inouts(modules, fsu, doc)
        # СОЗДАЕМ РАЗДЕЛ С ПАРАМЕТРИРОВАНИЕ СВЕТОДИОДОВ И ФК
        self._create_section_leds(modules, hmi, fsu, doc) 
        # СОЗДАЕМ РАЗДЕЛ С КОНФИГУРАЦИЕЙ
        self._create_section_config(aux_funcs, doc)
        # СОЗДАЕМ РАЗДЕЛ С ПАРАМЕТРАМИ РЕГИСТРАЦИИ
        self._create_section_disturb(fsu, doc)
        # ДОБАВЛЯЕМ ФИНАЛЬНУЮ ТАБЛИЦУ С ПОДПИСЯМИ
        add_new_section(doc)
        add_table_final(doc)
        doc.save("temp.docx")

    def get_blanc(self, device):
        self.create_template(device.fsu, device.modules, device.hmi, device.aux_funcs)

        doc = DocxTemplate('temp.docx')

        # Общие данные для титульных листов
        device_name = '-'.join(device.order_code.split('-')[:3])
        title = device.full_description + '\n' + f'«{device_name}»'

        colontile = ''
        if self.versions:
            last_version = self.versions[-1]
            colontile = f"Редакция {last_version['edition']} от {last_version['data']}"

        general_data = {"title":title, "code":self.code, "device_name": device_name, "device_order_code": device.order_code, "hmi_order_code": device.order_code_hmi if device.order_code_hmi!='' else 'Не заказан', "versions":self.versions[-10:], "colontile": colontile}

        context = {
            'fsu': device.fsu,
            'modules': device.modules,
            'hmi': device.hmi,
            'aux_funcs': device.aux_funcs,
            'general_data': general_data
        }

        name_for_save = f"{self.code} Бланк уставок {device_name} ред.{last_version['edition']}"

        doc.render(context)
        doc.save(f'{name_for_save}.docx')
        Logger.info(f"Бланк уставок сохранен в файл: '{name_for_save}.docx'")