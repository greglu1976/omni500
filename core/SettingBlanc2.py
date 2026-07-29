# Класс, представляющий бланк уставок
# Требует инициализации
# для генерации бланка уставок в него нужно передать объект device_data

import re
import json

from docxtpl import DocxTemplate
from docx import Document

from utils.docx_handler import add_new_section, add_new_section_landscape
from utils.tables import add_table_final, add_table_settings_core4, add_table_mtrx_ins_core4, add_table_mtrx_outs_core4, add_table_leds_new_core4, add_table_fks_core4, add_table_binaries_core4, add_table_reg_core4, add_table_settings_core4_simple

from xml.sax.saxutils import escape # для экранирования в дропдаун списке всяких << >>

from docxtpl import DocxTemplate

from logger.logger import Logger

from docx.shared import RGBColor # Добавьте этот импорт в начало файла

from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from core.Manual import Manual # Импорт вашего класса Manual
from core.FBData import FBData  # <--- ОБЯЗАТЕЛЬНО ДОБАВИТЬ

class SettingBlanc:
    def __init__(self, device_data):

        Logger.info(f"Путь к пакетам поддержки: {device_data['path_to_json']}")
        self.packet_path = device_data['path_to_json']
        self.device_data = device_data
        self.code = self.device_data["setting_blanc_code"]
        self.versions = self.device_data["versions"]
        self.base_structure = None  # Будет хранить структуру из get_all_settings()

        # Создаем экземпляр Manual для использования его движка данных
        self.manual_engine = Manual(device_data)

        

    def _create_section_settings_core4(self, doc, mode=2):
        """
        Генерирует раздел уставок.
        mode=6: Полный формат (№, Name, Description, Range, Unit, Step, Default)
        mode=5: Сокращенный формат (№, Description[as Name], Range, Unit, Step, Default) - аналог table5cols
        """

        add_new_section_landscape(doc)
        p = doc.add_paragraph('УСТАВКИ РЗиА')
        p.style = 'ДОК Заголовок 1'

        Logger.info(f"Запрос структурированных данных уставок через Manual (mode={mode})...")
        
        # Создаем помощник для форматирования чисел
        fb_helper = FBData({}) 
        
        # Получаем данные
        structured_data = self.manual_engine.get_all_settings_structured()
        
        if not structured_data:
            Logger.warning("Нет данных уставок от Manual.")
            return

        total_rows = 0

        for block in structured_data:
            for group in block["settings_groups"]:
                if group["MacroBlock"] != '-':
                     p_group = doc.add_paragraph(group["MacroBlock"])
                     p_group.style = 'ДОК Таблица Название'

                # Создаем таблицу. 
                # ВАЖНО: Убедитесь, что add_table_settings_core4 создает таблицу 
                # с правильным количеством заголовков для выбранного режима.
                # Если функция одна и та же, она должна создавать универсальную структуру.
                if mode == 1:
                    table = add_table_settings_core4_simple(doc)
                else:
                    table = add_table_settings_core4(doc)

                local_row_index = 1 # Сброс нумерации для каждой таблицы/группы

                for setting in group["Settings"]:
                    row = table.add_row()
                    total_rows += 1
                    
                    # --- ЗАПОЛНЕНИЕ ЯЧЕЕК В ЗАВИСИМОСТИ ОТ РЕЖИМА ---
                    
                    # 1. Номер (всегда в ячейке 0)
                    row.cells[0].text = str(local_row_index)
                    row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    if mode == 2:
                        # === РЕЖИМ 6 КОЛОНОК (ПОЛНЫЙ) ===
                        
                        # 2. Наименование (Name) - ячейка 1
                        self._set_cell_with_color(row.cells[1], setting["Name"])

                        # 3. Описание (Description) - ячейка 2
                        self._set_cell_with_color(row.cells[2], setting["Description"])

                        # Остальные данные
                        min_val = fb_helper._format_by_step(setting["Min"], setting["Step"]).replace('.', ',')
                        max_val = fb_helper._format_by_step(setting["Max"], setting["Step"]).replace('.', ',')
                        predefined = setting.get("PredefinedValues", "")
                        
                        if predefined:
                            range_text = predefined.replace('\\\\', ' / ')
                            default_text = self._get_default_from_enum_safe(setting["Default"], predefined)
                            unit_text = "-"
                            step_text = "-"
                        else:
                            range_text = f"{min_val} ... {max_val}"
                            default_val = fb_helper._format_by_step(setting["Default"], setting["Step"]).replace('.', ',')
                            default_text = default_val if default_val else ""
                            unit_text = setting.get("Unit", "-")
                            step_text = str(setting["Step"]).replace('.', ',')

                        # 4. Диапазон - ячейка 3
                        row.cells[3].text = range_text
                        row.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # 5. Ед. изм. - ячейка 4
                        row.cells[4].text = unit_text
                        row.cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # 6. Шаг - ячейка 5
                        row.cells[5].text = step_text
                        row.cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # 7. По умолчанию - ячейка 6
                        row.cells[6].text = default_text
                        row.cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    elif mode == 1:
                        # === РЕЖИМ 5 КОЛОНОК (СОКРАЩЕННЫЙ, как в LaTeX table5cols) ===
                        # Столбец "ПО ЮС" (Name) убираем. 
                        # В столбец "Наименование" (ячейка 1) ставим Description.
                        
                        # 2. Наименование (берем Description) - ячейка 1
                        self._set_cell_with_color(row.cells[1], setting["Description"])

                        # Остальные данные те же
                        min_val = fb_helper._format_by_step(setting["Min"], setting["Step"]).replace('.', ',')
                        max_val = fb_helper._format_by_step(setting["Max"], setting["Step"]).replace('.', ',')
                        predefined = setting.get("PredefinedValues", "")
                        
                        if predefined:
                            range_text = predefined.replace('\\\\', ' / ')
                            default_text = self._get_default_from_enum_safe(setting["Default"], predefined)
                            unit_text = "-"
                            step_text = "-"
                        else:
                            range_text = f"{min_val} ... {max_val}"
                            default_val = fb_helper._format_by_step(setting["Default"], setting["Step"]).replace('.', ',')
                            default_text = default_val if default_val else ""
                            unit_text = setting.get("Unit", "-")
                            step_text = str(setting["Step"]).replace('.', ',')

                        # 3. Диапазон - ячейка 2 (сдвиг на 1 влево)
                        row.cells[2].text = range_text
                        row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # 4. Ед. изм. - ячейка 3
                        row.cells[3].text = unit_text
                        row.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # 5. Шаг - ячейка 4
                        row.cells[4].text = step_text
                        row.cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # 6. По умолчанию - ячейка 5
                        row.cells[5].text = default_text
                        row.cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        # Если в таблице есть лишняя 6-я ячейка, очищаем её
                        if len(row.cells) > 6:
                            row.cells[6].text = ""

                    # Установка шрифта для всех ячеек строки
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                    
                    local_row_index += 1

        Logger.info(f"Сгенерировано {total_rows} строк уставок.")

    def _set_cell_with_color(self, cell, text):
        """Вспомогательный метод для установки текста с обработкой красного цвета"""
        if "\\textcolor{red}" in text:
            clean_text = text.replace("\\textcolor{red}{", "").rstrip("}")
            cell.text = ""
            run = cell.paragraphs[0].add_run(clean_text)
            run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            cell.text = text


    def _get_default_from_enum_safe(self, default_index, predefined_values_str):
        """Локальная копия или вызов метода из Manual"""
        try:
            if not isinstance(predefined_values_str, str) or not isinstance(default_index, int):
                return ""
            values = [v.strip() for v in predefined_values_str.split('\\\\')]
            if 0 <= default_index < len(values):
                return values[default_index]
        except:
            pass
        return ""


    def _fill_table_settings(self, table, rows_data):
        """
        Заполняет таблицу уставок данными для Core4
        
        Логика обработки col1 (ПО ЮС):
        1. Если есть текст в круглых скобках, например "Ввод функции в работу (Ввод_функции)",
        то последние скобки "(Ввод_функции)" переносятся в col2 (ИЧМ).
        2. Из col1 последние скобки и их содержимое удаляются.
        3. Если в строке несколько скобок, например "текст (первое) (второе)",
        то берется последнее вхождение: col1 = "текст (первое)", col2 = "второе"
        4. Старое значение col2 игнорируется/перезаписывается.
        """
        
        # Добавляем строки с данными (начиная с row_index=2, т.к. 0 и 1 - заголовки)
        for i, row_data in enumerate(rows_data, start=1):
            row = table.add_row()
              
            # --- Заполнение таблицы ---
            
            # Ячейка 0 (№) - номер по порядку
            row.cells[0].text = str(i)
            row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Ячейка 1 (ПО ЮС) - очищенное значение
            #row.cells[1].text = row_data[0]


            raw_text = row_data[0] # Или setting["Description"]
            
            # Проверяем, является ли это специальной строкой ошибки LaTeX
            if "\\textcolor{red}" in raw_text:
                # Очищаем текст от LaTeX тегов
                # Удаляем \textcolor{red}{ и последнюю }
                clean_text = raw_text.replace("\\textcolor{red}{", "").rstrip("}")
                
                cell = row.cells[1]
                cell.text = "" # Очищаем ячейку перед добавлением run
                run = cell.paragraphs[0].add_run(clean_text)
                run.font.color.rgb = RGBColor(255, 0, 0) # Красный цвет
            else:
                # Обычный текст
                row.cells[1].text = raw_text


            # Ячейка 2 (ИЧМ) - значение из последних скобок (или пусто)
            row.cells[2].text = row_data[1]
            row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Ячейка 3 (Значение / Диапазон) - из col3 с обработкой note_
            row.cells[3].text = str(row_data[2]) + " ... " + str(row_data[3])
            row.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
            
            # Ячейка 4 (Ед. изм.) - из col4
            row.cells[4].text = row_data[4]
            row.cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Ячейка 5 (Шаг) - из col5
            row.cells[5].text = row_data[5]
            row.cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Ячейка 6 (Значение по умолчанию) - из col6
            row.cells[6].text = row_data[6]
            row.cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Ячейки 7-10 (Группы уставок) - не заполняем
            row.cells[7].text = ''
            row.cells[8].text = ''
            row.cells[9].text = ''
            row.cells[10].text = ''
            
            # Устанавливаем размер шрифта
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        return table
        
    
    def _parse_note_dict(self, note_str):
        """
        Парсит строку вида "note_{'0': 'Не предусмотрено', '1': 'Предусмотрено'}"
        в формат "Не предусмотрено / Предусмотрено"
        """
        try:
            import ast
            # Убираем "note_" в начале
            if note_str.startswith('note_{'):
                dict_str = note_str[5:]  # оставляем "{'0': '...', '1': '...'}"
                note_dict = ast.literal_eval(dict_str)
                
                # Извлекаем только значения и сортируем их по ключам, 
                # чтобы порядок был предсказуемым (0, 1, 2...), так как словари в старых Python не упорядочены,
                # а в новых хотя и сохраняют порядок вставки, но явная сортировка надежнее для конфигов.
                # Если ключи всегда строковые цифры, можно отсортировать как числа или как строки.
                sorted_items = sorted(note_dict.items(), key=lambda item: int(item[0]) if item[0].isdigit() else item[0])
                values = [str(v) for k, v in sorted_items]
                
                return ' /\n'.join(values)
        except Exception as e:
            Logger.warning(f"Ошибка парсинга note_dict: {e}")
        
        return note_str




    def prepare_structure(self):
        """Подготовка полной структуры данных для LaTeX рендера"""
        
        # Создаем словарь для быстрого доступа к блокам по ID
        blocks_by_id = {info['Id']: info for info in self.fsu_information}
        
        # Находим все корневые блоки (ParentId == None)
        root_blocks = [info for info in self.fsu_information if info['ParentId'] is None]
        
        # Сортируем корневые блоки по имени
        root_blocks.sort(key=lambda x: x['Name'])
        
        latex_structure = []
        
        for root in root_blocks:
            # Информация о корневом блоке
            root_entry = {
                'type': 'root',
                'id': root['Id'],
                'name': root['Name'],
                'display_name': root['DisplayName'],
                'type_name': root['TypeName'],
                'variables': root.get('Variables', []),
                'children': [],
                'level': 0,
                'is_root': True
            }
            
            # Находим все дочерние блоки
            children = [info for info in self.fsu_information 
                    if info.get('ParentId') == root['Id']]
            children.sort(key=lambda x: x['Name'])
            
            for child in children:
                child_entry = {
                    'type': 'child',
                    'id': child['Id'],
                    'name': child['Name'],
                    'display_name': child['DisplayName'],
                    'type_name': child['TypeName'],
                    'variables': child.get('Variables', []),
                    'parent_id': child['ParentId'],
                    'parent_name': root['Name'],
                    'level': 1,
                    'is_root': False
                }
                root_entry['children'].append(child_entry)
            
            latex_structure.append(root_entry)
        
        return latex_structure




    def get_all_settings(self):
        """Собирает структуру уставок из заказа"""

        #self.maps = self.order_handler.get_mapping()
        
        ordered_fbs = list(self.maps.keys())
        base_structure = []
        
        for fb in ordered_fbs:
            fb_map = self.maps.get(fb)
            if not fb_map:
                continue
            
            json_data = self.order_handler.get_data_by_fb_name(fb_map)
            #print(json_data)
            parsed_blocks = self.order_handler.parse_rza_structure(json_data, all_struct=0)
            base_structure.extend(parsed_blocks)
        
        self.base_structure = base_structure
        Logger.info(f"Загружено {len(base_structure)} блоков уставок")
        return base_structure



    def create_template(self, mode):

        # Загрузка вспомогательного файла, где находятится полное описание
        with open(self.packet_path+'/fsu-information.json', 'r', encoding='utf-8') as f:
            self.fsu_information = json.load(f)['FunctionalBlocksInformation']

        # Загрузка вспомогательного файла, где находится полное описание
        with open(self.packet_path + '/meta.json', 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Собираем необходимые данные по устройству     
        DeviceSpecification = data['DeviceSpecification']
        HmiSpecification = data['HmiSpecification']
        # Сохраняем количество групп уставок
        self.SettingGroupCount = DeviceSpecification['SettingGroupCount']

        last_version = ""
        colontile = ''
        if self.device_data['versions']:
            last_version = self.device_data['versions'][-1]
            colontile = f"Редакция {last_version['edition']} от {last_version['data']}"

        def order_to_str(order_code):
            result = '-'.join(order_code.values())
            return result if result.strip() else "Не указано"
        
        first_part = order_to_str(DeviceSpecification["OrderCode"])
        second_part = order_to_str(HmiSpecification["OrderCode"])

        context = {
            "title": self.device_data['full_description'],
            "code": self.device_data['setting_blanc_code'],
            "device_order_code": first_part,
            "hmi_order_code": second_part,
            "versions":  self.device_data['versions'],
            "device_name":  self.device_data['name'],
            "colontile": colontile,
            "packet": data["Version"],
            "version": last_version['edition'],
            "date": last_version['data']
        }

        # Создаем документ
        doc_tpl = DocxTemplate('origin.docx')
        doc_tpl.render(context)
        doc_tpl.save('temp.docx')
        doc = Document('temp.docx')


        # Получаем структуру уставок
        #self.get_all_settings()


        self.base_structure = self.prepare_structure()
        # Генерируем раздел уставок (новый метод)
        Logger.info("Создаем раздел Уставки РЗиА...")

        self._create_section_settings_core4(doc, mode)




        Logger.info("Создаем раздел Матрица входов и выходных реле...")
        #self._create_section_inouts_core4(doc)

        #if second_part:
            #Logger.info("ИЧМ присутствует. Создаем раздел Настройка светодиодов и ФК...")
            #self._create_section_leds_core4(second_part, doc)
        Logger.info("Создаем раздел Конфигурация...")
        #self._create_section_config_core4(doc)
        Logger.info("Создаем раздел Натройка регистрации...")
        #self._create_section_disturb_core4(doc)
        # Остальные разделы пока закомментированы, при необходимости аналогично адаптировать
        # self._create_section_disturb_core4(device.fsu, doc)
        
        # Добавляем финальную таблицу
        add_new_section(doc)
        add_table_final(doc)
        
        # Сохраняем

        name_for_save = f"{self.code} Бланк уставок {self.device_data['name']} ред.{last_version['edition']}"
        #name_for_save = f"{self.code} Бланк уставок Core4"
        doc.save(f'{name_for_save}.docx')
        Logger.info(f"Бланк уставок сохранен: '{name_for_save}.docx'")
        
        return doc

    def get_blanc(self, mode):
        """
        Основной метод для генерации бланка уставок Core4
        """
        #print(device_data)
        self.create_template(mode)


##################################################################################
#####################################################################################
####################################################################################


    # РАЗДЕЛ ПАРАМЕТРИРОВАНИЯ ВХОДОВ И ВЫХОДОВ
    def _create_section_inouts_core4(self, doc):
        """
        Генерирует раздел документации "Матрица входов и выходов".
        """
        # ======================================================================
        # ЧАСТЬ 0: Подготовка общих списков сигналов для Dropdown
        # ======================================================================
        try:
            raw_sigs, raw_di_sigs = self.order_handler.get_fsu_signals()
        except Exception:
            raw_sigs, raw_di_sigs = [], []

        def extract_description(item):
            if isinstance(item, dict):
                return (item.get('fullDescription') or 
                        item.get('appliedDescription') or 
                        item.get('description') or 
                        item.get('name', ''))
            return str(item)

        # Очищенные списки строк для dropdown
        sigs_list = [desc for desc in [extract_description(s) for s in raw_sigs] if desc]
        di_sigs_list = [desc for desc in [extract_description(s) for s in raw_di_sigs] if desc]
        self.di_list =  di_sigs_list
        # Получаем данные слотов один раз
        slots_data = self.order_handler.get_slots_data()
        items_to_process = []
        if isinstance(slots_data, list):
            for slot_dict in slots_data:
                items_to_process.extend(slot_dict.items())
        elif isinstance(slots_data, dict):
            items_to_process.extend(slots_data.items())

        # ======================================================================
        # ЧАСТЬ 1: ОБРАБОТКА ВХОДОВ (M*_B*_B*_Status)
        # ======================================================================
        pattern_inputs = re.compile(r'^M\d+_B\d{3}_B\d+_Status$')
        final_dict_inputs = {}
        
        for slot_name, params_list in items_to_process:
            status_signals = [p for p in params_list if pattern_inputs.match(p)]
            if not status_signals: continue
            
            clean_sig_list = []
            for sig in status_signals:
                try:
                    d = self.config_handler.get_param_info(sig)
                    desc = d.get("appliedDescription", sig)
                    clean_sig_list.append(desc.replace(". Статус", "").strip())
                except:
                    clean_sig_list.append(sig)
            if clean_sig_list:
                final_dict_inputs[slot_name] = clean_sig_list

        # ======================================================================
        # ЧАСТЬ 2: ОБРАБОТКА ВЫХОДОВ (M*_K*_K*_Status)
        # ======================================================================
        pattern_outputs = re.compile(r'^M\d+_K\d{3}_K\d+_Status$')
        final_dict_outputs = {}
        
        sigs_list_outputs = self.order_handler.get_fsu_out_signals()
        
        outs_list = self.extract_all_signals_from_structure(sigs_list_outputs)

        for slot_name, params_list in items_to_process:
            #print(slot_name, params_list)
            status_signals = [p for p in params_list if pattern_outputs.match(p)]
            if not status_signals: continue
            
            clean_sig_list = []
            for sig in status_signals:
                try:
                    d = self.config_handler.get_param_info(sig)
                    desc = d.get("appliedDescription", sig)
                    clean_sig_list.append(desc.replace(". Статус", "").strip())
                except:
                    clean_sig_list.append(sig)
            if clean_sig_list:
                final_dict_outputs[slot_name] = clean_sig_list

        # ======================================================================
        # ЧАСТЬ 3: ГЕНЕРАЦИЯ ДОКУМЕНТА
        # ======================================================================
        if not final_dict_inputs and not final_dict_outputs:
            return

        add_new_section_landscape(doc) 
        
        p = doc.add_paragraph('МАТРИЦА ВХОДОВ И ВЫХОДНЫХ РЕЛЕ')
        p.style = 'ДОК Заголовок 1'

        # --- ГЕНЕРАЦИЯ ВХОДОВ ---
        if final_dict_inputs:
            doc.add_paragraph('Дискретные входы').style = 'ДОК Заголовок 2'
            
            sorted_inputs = sorted(final_dict_inputs.keys(), key=lambda x: int(re.search(r'M(\d+)', x).group(1)) if re.search(r'M(\d+)', x) else 0)
            
            for slot_name in sorted_inputs:
                doc.add_paragraph(f"{slot_name}").style = 'ДОК Таблица Название'
                add_table_mtrx_ins_core4(
                    doc=doc,
                    slot_name=slot_name,
                    inputs_list=final_dict_inputs[slot_name],
                    sigs=sigs_list,
                    di_sigs=di_sigs_list
                )
                doc.add_paragraph()

        # --- ГЕНЕРАЦИЯ ВЫХОДОВ ---
        if final_dict_outputs:
            #print(final_dict_outputs)
            doc.add_paragraph('Выходные реле').style = 'ДОК Заголовок 2'
            
            sorted_outputs = sorted(final_dict_outputs.keys(), key=lambda x: int(re.search(r'M(\d+)', x).group(1)) if re.search(r'M(\d+)', x) else 0)
            
            for slot_name in sorted_outputs:
                doc.add_paragraph(f"{slot_name}").style = 'ДОК Таблица Название'
                # Вызываем новую функцию для выходов

                #print(sigs_list)
                add_table_mtrx_outs_core4(
                    doc=doc,
                    outputs_list=final_dict_outputs[slot_name],
                    sigs_list=outs_list
                )
                doc.add_paragraph()

        return


    def extract_all_signals_from_structure(self, data_structure):
        """
        Извлекает все сигналы из структуры данных функций и подфункций.
        
        :param data_structure: Список словарей с ключами 'function' и 'subfunctions'
        :return: Плоский список всех сигналов (desc)
        """
        all_signals = []
        
        for function_item in data_structure:
            # Получаем название функции (для контекста, если нужно)
            function_name = function_item.get('function', '')
            
            # Проходим по всем подфункциям
            subfunctions = function_item.get('subfunctions', [])
            for subfunc in subfunctions:
                subfunc_name = subfunc.get('name', '')
                
                # Извлекаем данные (сигналы) из подфункции
                data_list = subfunc.get('data', [])
                for signal in data_list:
                    if signal:  # Пропускаем пустые строки
                        all_signals.append(signal)
        
        return all_signals


    # РАЗДЕЛ СВЕТОДИОДОВ И ФК
    def _create_section_leds_core4(self, order_code, doc):

        # парсим код заказа ИЧМ
        parts = order_code.split('-')
        result = [
            "Модуль расширения 1 на 16 светодиодов" if parts[3]=="С" else "Модуль расширения 1 на 16 функциональных кнопок" if parts[3]=="К" else "Модуль отсутствует",
            "Модуль расширения 2 на 16 светодиодов" if parts[4]=="С" else "Модуль расширения 2 на 16 функциональных кнопок" if parts[4]=="К" else "Модуль отсутствует",
            "Модуль расширения 3 на 16 светодиодов" if parts[5]=="С" else "Модуль расширения 3 на 16 функциональных кнопок" if parts[5]=="К" else "Модуль отсутствует",
            "Модуль расширения 4 на 16 светодиодов" if parts[6]=="С" else "Модуль расширения 4 на 16 функциональных кнопок" if parts[6]=="К" else "Модуль отсутствует"
        ]

        #############################################################################
        # СОЗДАЕМ РАЗДЕЛ НАСТРОЙКА СВЕТОДИОДОВ И ФУНКЦИОНАЛЬНЫХ КЛАВИШ
        add_new_section_landscape(doc) # Создаем раздел для матрицы вх/вых
        # Добавляем заголовок
        p = doc.add_paragraph('НАСТРОЙКА СВЕТОДИОДОВ И ФУНКЦИОНАЛЬНЫХ КЛАВИШ')
        p.style = 'ДОК Заголовок 1'
        p = doc.add_paragraph('Светодиоды')
        p.style = 'ДОК Заголовок 2'

        # вытаскиваем выпадающий список сигналов для светодиодов
        drop_list = self.order_handler.get_digital_signals_for_led()

        doc.add_paragraph("ИЧМ").style = 'ДОК Таблица Название'
        add_table_leds_new_core4(doc, drop_list)

        for res in result:
            if "светодиодов" in res:
                doc.add_paragraph()
                doc.add_paragraph(res).style = 'ДОК Таблица Название'
                add_table_leds_new_core4(doc, drop_list)                

        ###############################################################
        p = doc.add_paragraph('Функциональные клавиши')
        p.style = 'ДОК Заголовок 2'

        doc.add_paragraph("ИЧМ").style = 'ДОК Таблица Название'        
        add_table_fks_core4(doc, self.di_list)

        for res in result:
            if "функциональных кнопок" in res:
                doc.add_paragraph()
                doc.add_paragraph(res).style = 'ДОК Таблица Название'
                add_table_fks_core4(doc, drop_list)  


    #########################################
    #################### РАЗДЕЛ КОНФИГУРАЦИЯ 
    #########################################

    def enum_calc(self, enum, default):
        result_str = " / ".join([item['VisibleValue'] for item in enum])
        default_visible = None
        for item in enum:
            if str(item.get('ParameterValue')) == default:
                default_visible = item.get('VisibleValue')
                break
        return result_str, default_visible

    def parse_note(self, note, default):
        """
        Парсит строку вида "1 - Вывод, 2 - Ввод, 3 - Неизвестно"
        """
        # Разбиваем строку на части
        parts = [part.strip() for part in note.split(',')]
        
        # Извлекаем значения
        values = []
        default_value = None
        
        for part in parts:
            if ' - ' in part:
                key_str, value = part.split(' - ', 1)
                key_str = key_str.strip()
                value = value.strip()
                values.append(value)
                
                # Сравниваем как строки (не преобразуем в int)
                if key_str == str(default):  # Приводим default к строке
                    default_value = value
        
        # Формируем строку для отображения всех вариантов
        result_str = " / ".join(values)
        
        return result_str, default_value

    def _create_section_config_core4(self, doc):

        add_new_section(doc)
        p = doc.add_paragraph('КОНФИГУРАЦИЯ')
        p.style = 'ДОК Заголовок 1'

        raw_data = self.order_handler.get_data_for_configuration()



        for datum in raw_data:
            if datum["main_title"] == "ИЧМ":
                continue
            p = doc.add_paragraph(datum["main_title"])
            p.style = 'ДОК Заголовок 2'

            for table in datum["tables"]:
                doc.add_paragraph(table["title"]).style = 'ДОК Таблица Название'

                fixed_rows = []
                for row in table["rows"]:

                    row_name = row["name"]
                    row_data = self.config_handler.get_param_info(row_name)
                    enum_data = self.extension_handler.find_enum_by_parameter_name(row_name)
                    #print(enum_data, row_name)

                    col1 = row_data["fullDescription"]
                    col2 = row_data["appliedDescription"]
                    col3 = row_data["note"] if row_data["note"].count('-') >= 2 else ""
                    col4 = row_data["units"] if row_data["units"] else '-'
                    col5 = row_data["step"] if row_data["step"] else '-'
                    col6 = row_data["defaultValue"]

                    if col3:
                        col3, col6 = self.parse_note(col3, col6)

                    if str(col6)=="false":
                        col3 = 'Вывод / Ввод'
                        col6 = "Вывод"

                    if str(col6)=="true":
                        col3 = 'Вывод / Ввод'
                        col6 = "Ввод"

                    if enum_data:
                        col3, col6 =  self.enum_calc(enum_data, col6)
                        col4 = col5 = '-'
                    elif col3:
                        col4 = col5 = '-'                           
                    elif row_data["minValue"] is None:
                        col3 = '-'
                    else:
                        if row_data["minValue"]=="0" and row_data["maxValue"]=="1":
                            col3 = 'Вывод / Ввод'
                            col6 = "Вывод" if row_data["defaultValue"]=="0" else "Ввод"
                            col5 = '-'
                        else:    
                            col3 = row_data["minValue"] + ' ... ' + row_data["maxValue"]

                    fixed_rows.append((col1, col2, col3, col4, col5, col6))
                add_table_binaries_core4(doc, fixed_rows)
                p = doc.add_paragraph()


    def _create_section_disturb_core4(self, doc):
        
        add_new_section_landscape(doc)

        p = doc.add_paragraph('НАСТРОЙКА РЕГИСТРАЦИИ')
        p.style = 'ДОК Заголовок 1'

        reg_data = self.order_handler.get_data_for_registration()

        for section in reg_data:
            p = doc.add_paragraph(section["main_title"])
            p.style = 'ДОК Заголовок 2'
            
            # Обрабатываем подразделы
            for subsection in section.get("subsections", []):
                # Заголовок подраздела (например, "ТО РПН")
                _name = subsection["title"].split("_")[0]
                p_sub = doc.add_paragraph(self.abbr_dict.get(_name, _name))    #)  subsection["title"])
                p_sub.style = 'ДОК Заголовок 3'  # или другой стиль для подраздела
                
                # Таблицы внутри подраздела
                for table_data in subsection["tables"]:
                    doc.add_paragraph(table_data["title"]).style = 'ДОК Таблица Название'
                    
                    data_rows = []
                    for param_name in table_data["parameters"]:
                        row_info = self.config_handler.get_param_info(param_name)
                        
                        if row_info:
                            col1 = row_info.get("fullDescription", "")
                            col2 = row_info.get("appliedDescription", "")
                            param_type = row_info.get("type")
                            
                            data_rows.append((col1, col2, param_type))
                    
                    if data_rows:
                        add_table_reg_core4(doc, data_rows)
                        doc.add_paragraph().style = 'TAGS'


    def get_table_settings_latex1(self, ln, fb):

        #print(ln, fb)
        if not self.base_structure:
            self.get_all_settings()
        #print(self.base_structure)

        for bloc in self.base_structure:
            if bloc["type"]=="simple":
                a = bloc["rows"][0]["col0"].split("_1_")
                if a[0]==fb: #and a[1].split('_')[0]==ln:
                    return bloc
            else:
                subs = bloc["sub_functions"]
                for sub in subs:
                    a = sub["rows"][0]["col0"].split("_1_")
                    if a[0]==fb: #and a[1].split('_')[0]==ln:
                        #return sub
                        return bloc
        return None
    

    def get_table_settings_latex(self, ln, fb):

        #print(ln, fb)
        if not self.base_structure:
            self.get_all_settings()
        #print(self.base_structure)

        for bloc in self.base_structure:
            if bloc["type"] == "simple":
                col0_value = bloc["rows"][0]["col0"]
                
                # Сначала пробуем _1_
                if "_1_" in col0_value:
                    a = col0_value.split("_1_")
                # Если нет _1_, пробуем _2_
                elif "_2_" in col0_value:
                    a = col0_value.split("_2_")
                    Logger.error(f"Разделитель _2_ !!! {col0_value}")
                else:
                    continue
                    
                if a[0] == fb:  # and a[1].split('_')[0] == ln:
                    return bloc
            else:
                subs = bloc["sub_functions"]
                for sub in subs:
                    col0_value = sub["rows"][0]["col0"]
                    
                    # Сначала пробуем _1_
                    if "_1_" in col0_value:
                        a = col0_value.split("_1_")
                    # Если нет _1_, пробуем _2_
                    elif "_2_" in col0_value:
                        a = col0_value.split("_2_")
                        Logger.error(f"Разделитель _2_ !!! {col0_value}")
                    else:
                        continue
                        
                    if a[0] == fb:  # and a[1].split('_')[0] == ln:
                        #return sub
                        return bloc
        return None
